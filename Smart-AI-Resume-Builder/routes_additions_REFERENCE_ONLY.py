"""
==================================================================================
  ROUTES ADDITIONS — paste these route functions into your existing Flask app.py
==================================================================================
  Missing features restored:
    1. Resume Builder  (/builder, /builder POST)
    2. AI Interview Questions  (/ai-questions, /ai-questions POST)
    3. Job Search  (/job-search)
    4. Feedback  (/feedback, /feedback POST)
    5. About  (/about)
==================================================================================
"""

# ── Add these imports at the top of your app.py ──────────────────────────────
#
#   import re, random, json, io
#   import pandas as pd
#   from docx import Document
#   from docx.shared import Pt, Inches
#   from docx.enum.text import WD_ALIGN_PARAGRAPH
#   from flask import send_file, request, session
#   from datetime import datetime
#
# ── Make sure these exist in your project (copy from Streamlit project) ───────
#   csv/Properly_Categorized_Interview_Questions.csv
#   (the ResumeBuilder utility can be ported, or use the inline DOCX builder below)
#
# ─────────────────────────────────────────────────────────────────────────────


# ══════════════════════════════════════════════════════════════════════════════
#  RESUME BUILDER
# ══════════════════════════════════════════════════════════════════════════════

@app.route('/builder', methods=['GET'])
@login_required
def builder():
    if current_user.role != 'user':
        return redirect(url_for('recruiter'))
    return render_template('builder.html')


@app.route('/builder', methods=['POST'])
@login_required
def builder_post():
    """
    Receives the multi-section builder form, builds a DOCX resume,
    and streams it back as a file download.
    """
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import io

    data = request.form

    # ── Personal Info ─────────────────────────────────────────────────────
    full_name   = data.get('full_name', '').strip()
    email       = data.get('email', '').strip()
    phone       = data.get('phone', '').strip()
    location    = data.get('location', '').strip()
    linkedin    = data.get('linkedin', '').strip()
    portfolio   = data.get('portfolio', '').strip()
    summary     = data.get('summary', '').strip()
    template    = data.get('template', 'Modern')

    if not full_name or not email:
        flash('Full name and email are required.', 'error')
        return redirect(url_for('builder'))

    # ── Parse repeated sections (sent as JSON strings from JS) ────────────
    import json

    def parse_json_field(key):
        raw = data.get(key, '[]')
        try:
            return json.loads(raw)
        except Exception:
            return []

    experiences  = parse_json_field('experiences_json')
    educations   = parse_json_field('education_json')
    projects     = parse_json_field('projects_json')
    skills_data  = parse_json_field('skills_json')   # {technical,soft,languages,tools}

    # ── Build DOCX ────────────────────────────────────────────────────────
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin   = Inches(1.0)
        section.right_margin  = Inches(1.0)

    # Style helpers
    def add_heading(text, level=1):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text.upper())
        run.bold = True
        run.font.size = Pt(11) if level == 2 else Pt(13)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after  = Pt(2)
        p.add_run()  # separator line via border not available without XML; use rule
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '2196F3')
        pBdr.append(bottom)
        pPr.append(pBdr)

    def add_bullet(text):
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(text).font.size = Pt(10)
        p.paragraph_format.space_after = Pt(1)

    def body_text(text):
        p = doc.add_paragraph(text)
        p.runs[0].font.size = Pt(10) if p.runs else None
        p.paragraph_format.space_after = Pt(2)

    # ── Name & Contact ─────────────────────────────────────────────────────
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_para.add_run(full_name)
    name_run.bold = True
    name_run.font.size = Pt(20)

    contact_parts = [x for x in [email, phone, location, linkedin, portfolio] if x]
    contact_para = doc.add_paragraph(' | '.join(contact_parts))
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_para.runs[0].font.size = Pt(9)
    contact_para.paragraph_format.space_after = Pt(4)

    # ── Summary ────────────────────────────────────────────────────────────
    if summary:
        add_heading('Professional Summary', 2)
        body_text(summary)

    # ── Experience ─────────────────────────────────────────────────────────
    if experiences:
        add_heading('Work Experience', 2)
        for exp in experiences:
            p = doc.add_paragraph()
            p.add_run(exp.get('position', '')).bold = True
            p.add_run(f"  —  {exp.get('company', '')}")
            date_str = f"{exp.get('start_date', '')} – {exp.get('end_date', 'Present')}"
            tab_run = p.add_run(f"\t{date_str}")
            tab_run.font.size = Pt(9)
            p.paragraph_format.space_after = Pt(1)
            if exp.get('description'):
                body_text(exp['description'])
            for r in exp.get('responsibilities', []):
                if r:
                    add_bullet(r)
            for a in exp.get('achievements', []):
                if a:
                    add_bullet(f"✓ {a}")

    # ── Education ──────────────────────────────────────────────────────────
    if educations:
        add_heading('Education', 2)
        for edu in educations:
            p = doc.add_paragraph()
            p.add_run(f"{edu.get('degree', '')} in {edu.get('field', '')}").bold = True
            p.add_run(f"  —  {edu.get('school', '')}")
            grad = edu.get('graduation_date', '')
            gpa  = edu.get('gpa', '')
            meta = '  |  '.join([x for x in [grad, f"GPA: {gpa}" if gpa else ''] if x])
            if meta:
                p.add_run(f"\t{meta}").font.size = Pt(9)
            for a in edu.get('achievements', []):
                if a:
                    add_bullet(a)

    # ── Projects ───────────────────────────────────────────────────────────
    if projects:
        add_heading('Projects', 2)
        for proj in projects:
            p = doc.add_paragraph()
            p.add_run(proj.get('name', '')).bold = True
            tech = proj.get('technologies', '')
            if tech:
                p.add_run(f"  |  {tech}").font.size = Pt(9)
            if proj.get('link'):
                p.add_run(f"  [{proj['link']}]").font.size = Pt(9)
            if proj.get('description'):
                body_text(proj['description'])
            for r in proj.get('responsibilities', []):
                if r:
                    add_bullet(r)
            for a in proj.get('achievements', []):
                if a:
                    add_bullet(f"✓ {a}")

    # ── Skills ─────────────────────────────────────────────────────────────
    if skills_data:
        add_heading('Skills', 2)
        skill_lines = []
        if skills_data.get('technical'):
            skill_lines.append(('Technical', ', '.join(skills_data['technical'])))
        if skills_data.get('soft'):
            skill_lines.append(('Soft Skills', ', '.join(skills_data['soft'])))
        if skills_data.get('languages'):
            skill_lines.append(('Languages', ', '.join(skills_data['languages'])))
        if skills_data.get('tools'):
            skill_lines.append(('Tools & Tech', ', '.join(skills_data['tools'])))
        for label, val in skill_lines:
            p = doc.add_paragraph()
            p.add_run(f"{label}: ").bold = True
            p.add_run(val).font.size = Pt(10)
            p.paragraph_format.space_after = Pt(1)

    # ── Stream as download ─────────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    filename = f"{full_name.replace(' ', '_')}_resume.docx"
    return send_file(
        buf,
        as_attachment=True,
        download_name=filename,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )


# ══════════════════════════════════════════════════════════════════════════════
#  AI INTERVIEW QUESTIONS
# ══════════════════════════════════════════════════════════════════════════════

@app.route('/ai-questions', methods=['GET', 'POST'])
@login_required
def ai_questions():
    if current_user.role != 'user':
        return redirect(url_for('recruiter'))

    questions   = []
    prompt      = ''
    skills_str  = ''
    error       = None
    warning     = None
    encouraging = None

    ENCOURAGING = [
        "All the best for your interview!",
        "Hope these help you prepare well!",
        "Good luck with your interview preparation!",
        "You're going to ace this!",
        "Keep up the great work!",
        "Believe in yourself and your abilities!",
        "Remember, preparation is key to success!",
        "You've got this! Go show them what you're made of!",
        "Stay confident and positive!",
        "Every question is a step towards your goal!",
        "Practice makes perfect!",
        "You're one step closer to your dream job!",
        "Keep pushing forward!",
        "Success is just around the corner!",
        "Your hard work will pay off!",
    ]

    if request.method == 'POST':
        prompt = request.form.get('prompt', '').strip()

        if not prompt:
            warning = "Please enter a prompt to generate questions."
        else:
            # Load the CSV
            csv_path = os.path.join(os.path.dirname(__file__), 'csv',
                                    'Properly_Categorized_Interview_Questions.csv')
            try:
                import pandas as pd
                df = pd.read_csv(csv_path)
                if 'Question' not in df.columns or 'Skills' not in df.columns:
                    error = "CSV file must contain 'Question' and 'Skills' columns."
                else:
                    df['Skills'] = df['Skills'].apply(
                        lambda x: [s.strip().lower() for s in str(x).split(',')]
                    )

                    # Parse number of questions
                    match = re.search(r'(\d+)', prompt)
                    num_questions = min(int(match.group(1)), 50) if match else 50

                    # Parse skills from prompt
                    skills_match = re.search(r'on (.+)', prompt, re.IGNORECASE)
                    if not skills_match:
                        warning = "Please specify skills in your prompt (e.g., 'Give me 10 questions on Python and Java')."
                    else:
                        skills_text = skills_match.group(1)
                        user_skills = [
                            s.strip().lower()
                            for s in re.split(r',\s*|\s+and\s+', skills_text)
                            if s.strip()
                        ]

                        if not user_skills:
                            warning = "No skills specified. Please include skills in your prompt."
                        else:
                            filtered = df[df['Skills'].apply(
                                lambda x: any(sk in x for sk in user_skills)
                            )]

                            if filtered.empty:
                                warning = f"No questions found for: {', '.join(user_skills)}."
                            else:
                                if len(filtered) < num_questions:
                                    warning = f"Only {len(filtered)} questions available for the selected skills."
                                    questions = filtered['Question'].tolist()
                                else:
                                    questions = filtered.sample(n=num_questions)['Question'].tolist()

                                skills_str  = ', '.join(user_skills)
                                encouraging = random.choice(ENCOURAGING)

            except FileNotFoundError:
                error = ("Questions file not found. "
                         "Please ensure 'csv/Properly_Categorized_Interview_Questions.csv' exists.")
            except Exception as e:
                error = f"Error loading questions: {str(e)}"

    return render_template(
        'ai_questions.html',
        prompt=prompt,
        questions=questions,
        skills_str=skills_str,
        error=error,
        warning=warning,
        encouraging=encouraging,
    )


# ══════════════════════════════════════════════════════════════════════════════
#  JOB SEARCH
# ══════════════════════════════════════════════════════════════════════════════

@app.route('/job-search', methods=['GET', 'POST'])
@login_required
def job_search():
    if current_user.role != 'user':
        return redirect(url_for('recruiter'))

    results = []
    query   = ''
    location_q = ''
    error   = None

    if request.method == 'POST':
        query      = request.form.get('query', '').strip()
        location_q = request.form.get('location', '').strip()

        if not query:
            error = "Please enter a job title or keyword."
        else:
            # ── Adzuna API (free tier — set keys in env vars) ──────────────
            #   ADZUNA_APP_ID  and  ADZUNA_APP_KEY
            app_id  = os.environ.get('ADZUNA_APP_ID', '')
            app_key = os.environ.get('ADZUNA_APP_KEY', '')

            if app_id and app_key:
                try:
                    import requests as req_lib
                    country = 'in'   # change to 'gb', 'us', etc.
                    url = (f"https://api.adzuna.com/v1/api/jobs/{country}/search/1"
                           f"?app_id={app_id}&app_key={app_key}"
                           f"&results_per_page=20&what={req_lib.utils.quote(query)}"
                           f"&where={req_lib.utils.quote(location_q)}&content-type=application/json")
                    resp = req_lib.get(url, timeout=8)
                    if resp.ok:
                        data = resp.json()
                        for job in data.get('results', []):
                            results.append({
                                'title':    job.get('title', ''),
                                'company':  job.get('company', {}).get('display_name', 'N/A'),
                                'location': job.get('location', {}).get('display_name', 'N/A'),
                                'salary':   _fmt_salary(job),
                                'url':      job.get('redirect_url', '#'),
                                'desc':     job.get('description', '')[:200] + '…',
                                'date':     job.get('created', '')[:10],
                            })
                    else:
                        error = f"Job API error ({resp.status_code}). Please try again later."
                except Exception as e:
                    error = f"Could not fetch jobs: {str(e)}"
            else:
                # No API keys — show a helpful placeholder
                error = ("Job search requires Adzuna API credentials. "
                         "Set ADZUNA_APP_ID and ADZUNA_APP_KEY environment variables, "
                         "or visit adzuna.com to get free API keys.")

    return render_template(
        'job_search.html',
        results=results,
        query=query,
        location_q=location_q,
        error=error,
    )


def _fmt_salary(job):
    lo = job.get('salary_min')
    hi = job.get('salary_max')
    if lo and hi:
        return f"₹{int(lo):,} – ₹{int(hi):,}"
    if lo:
        return f"from ₹{int(lo):,}"
    return "Not specified"


# ══════════════════════════════════════════════════════════════════════════════
#  FEEDBACK
# ══════════════════════════════════════════════════════════════════════════════

# In-memory store — replace with a DB model in production
_feedback_store = []

@app.route('/feedback', methods=['GET'])
@login_required
def feedback():
    avg_rating = 0
    if _feedback_store:
        avg_rating = round(sum(f['rating'] for f in _feedback_store) / len(_feedback_store), 1)

    # Rating distribution
    dist = {1: 0, 2: 0, 3: 0, 4: 0, 5: 0}
    for f in _feedback_store:
        dist[f['rating']] = dist.get(f['rating'], 0) + 1

    return render_template(
        'feedback.html',
        feedbacks=_feedback_store,
        avg_rating=avg_rating,
        rating_dist=dist,
        total=len(_feedback_store),
    )


@app.route('/feedback', methods=['POST'])
@login_required
def feedback_post():
    name    = request.form.get('name', current_user.full_name).strip()
    email   = request.form.get('feedback_email', current_user.email).strip()
    rating  = int(request.form.get('rating', 3))
    feature = request.form.get('feature', '').strip()
    comment = request.form.get('comment', '').strip()

    if not comment:
        flash('Please write a comment before submitting.', 'error')
        return redirect(url_for('feedback'))

    _feedback_store.append({
        'name':    name,
        'email':   email,
        'rating':  rating,
        'feature': feature,
        'comment': comment,
        'date':    datetime.utcnow().strftime('%d %b %Y'),
    })

    flash('Thank you for your feedback! 🙏', 'success')
    return redirect(url_for('feedback'))


# ══════════════════════════════════════════════════════════════════════════════
#  ABOUT  (public page — no login required)
# ══════════════════════════════════════════════════════════════════════════════

@app.route('/about')
def about():
    return render_template('about.html')
