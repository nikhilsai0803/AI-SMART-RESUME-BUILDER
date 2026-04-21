import os, io, re, json, math, random, logging, joblib
from datetime import datetime
from functools import wraps
from concurrent.futures import ThreadPoolExecutor
from flask import (Flask, render_template, request, redirect,
                   url_for, session, flash, jsonify, Response, send_file)
from werkzeug.utils import secure_filename
from werkzeug.security import generate_password_hash, check_password_hash
from utils.resume_parser import ResumeParser
from utils.resume_analyzer import ResumeAnalyzer
from config.job_roles import JOB_ROLES

# ── Optional dependencies ────────────────────────────────────────────────────
try:
    import spacy; nlp = spacy.load("en_core_web_sm"); SPACY_LOADED = True
except Exception:
    nlp = None; SPACY_LOADED = False
try:
    from pdfminer.high_level import extract_text as pdfminer_extract; PDFMINER_LOADED = True
except Exception:
    PDFMINER_LOADED = False
try:
    from docx import Document as DocxDocument
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_LOADED = True
except Exception:
    DOCX_LOADED = False
try:
    import pandas as pd; PANDAS_LOADED = True
except Exception:
    PANDAS_LOADED = False

# ── App setup ────────────────────────────────────────────────────────────────
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

app = Flask(__name__)
app.secret_key = os.environ.get("FLASK_SECRET_KEY", "resumelens-ai-secret-2025")

from datetime import timedelta
app.config["PERMANENT_SESSION_LIFETIME"] = timedelta(days=30)

UPLOAD_FOLDER = os.path.join(BASE_DIR, "uploads")
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER
app.config["MAX_CONTENT_LENGTH"] = 10 * 1024 * 1024

# ── ML Models ────────────────────────────────────────────────────────────────
try:
    ats_model  = joblib.load(os.path.join(BASE_DIR, "ats_model.pkl"))
    vectorizer = joblib.load(os.path.join(BASE_DIR, "tfidf_vectorizer.pkl"))
    MODELS_LOADED = True
except Exception as e:
    ats_model = vectorizer = None; MODELS_LOADED = False
    logger.warning(f"ATS model not loaded: {e}")

parser   = ResumeParser()
analyzer = ResumeAnalyzer()

# ── In-memory stores (replace with DB models in production) ──────────────────
# Users:    { email: { id, full_name, email, password_hash, role, github, linkedin, created_at } }
# Feedback: [ { name, email, rating, feature, comment, date }, … ]
USERS           = {}
_feedback_store = []

# ── Seed two demo accounts on startup ────────────────────────────────────────
def _seed_demo():
    for email, name, role in [
        ("user@demo.com",      "Alex Johnson",   "user"),
        ("recruiter@demo.com", "Sarah Mitchell", "recruiter"),
    ]:
        if email not in USERS:
            USERS[email] = {
                "id":            email,
                "full_name":     name,
                "email":         email,
                "password_hash": generate_password_hash("password"),
                "role":          role,
                "github":        "",
                "linkedin":      "",
                "created_at":    datetime.utcnow(),
            }

_seed_demo()

# ════════════════════════════════════════════════════════════════════════════
# HELPERS & SHARED UTILITIES
# ════════════════════════════════════════════════════════════════════════════

ALLOWED_EXT = {"pdf", "docx"}

def allowed_file(fn):
    return "." in fn and fn.rsplit(".", 1)[1].lower() in ALLOWED_EXT

def score_color(s):
    s = float(s)
    return "success" if s >= 80 else ("warning" if s >= 60 else "danger")

def get_ats_ml(text):
    if MODELS_LOADED:
        try:
            return round(float(ats_model.predict(vectorizer.transform([text]))[0]), 1)
        except Exception:
            pass
    return None

def get_greeting():
    h = datetime.now().hour
    if h < 12: return "morning"
    if h < 17: return "afternoon"
    return "evening"

def current_user():
    """Return user dict from session, or None."""
    uid = session.get("user_id")
    if not uid:
        return None
    return USERS.get(uid)

def _fmt_salary(job):
    lo = job.get("salary_min")
    hi = job.get("salary_max")
    if lo and hi:
        return f"₹{int(lo):,} – ₹{int(hi):,}"
    if lo:
        return f"from ₹{int(lo):,}"
    return "Not specified"

# ── Jinja filters & globals ──────────────────────────────────────────────────
app.jinja_env.filters["score_color"] = score_color
app.jinja_env.filters["cos_rad"] = lambda deg: math.cos(math.radians(float(deg)))
app.jinja_env.filters["sin_rad"] = lambda deg: math.sin(math.radians(float(deg)))

@app.context_processor
def inject_user():
    """Make current_user available in every template as a simple object."""
    u = current_user()
    if u:
        class _U:
            is_authenticated = True
            def __init__(self, d):
                self.__dict__.update(d)
                self.full_name = d.get("full_name", "")
                self.role      = d.get("role", "user")
                self.email     = d.get("email", "")
                self.github    = d.get("github", "")
                self.linkedin  = d.get("linkedin", "")
        return {"current_user": _U(u)}
    class _Anon:
        is_authenticated = False
        full_name = ""
        role      = ""
        email     = ""
        github    = ""
        linkedin  = ""
    return {"current_user": _Anon()}

# ── Auth decorators ──────────────────────────────────────────────────────────
def login_required(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        if not session.get("user_id"):
            flash("Please sign in to continue.", "warning")
            return redirect(url_for("login", next=request.path))
        u = current_user()
        if u is None:
            # Session has stale user_id (e.g. app restarted) — clear and re-login
            session.clear()
            flash("Session expired. Please sign in again.", "warning")
            return redirect(url_for("login", next=request.path))
        return f(*args, **kwargs)
    return decorated

def recruiter_required(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        if not session.get("user_id"):
            flash("Please sign in.", "warning")
            return redirect(url_for("login"))
        u = current_user()
        if u is None:
            session.clear()
            flash("Session expired. Please sign in again.", "warning")
            return redirect(url_for("login"))
        if u.get("role") != "recruiter":
            flash("Recruiter access required.", "error")
            return redirect(url_for("dashboard"))
        return f(*args, **kwargs)
    return decorated

def user_required(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        if not session.get("user_id"):
            flash("Please sign in.", "warning")
            return redirect(url_for("login", next=request.path))
        u = current_user()
        if u is None:
            session.clear()
            flash("Session expired. Please sign in again.", "warning")
            return redirect(url_for("login"))
        if u.get("role") != "user":
            flash("This area is for job seekers.", "error")
            return redirect(url_for("recruiter"))  # recruiters go to their own dashboard
        return f(*args, **kwargs)
    return decorated

# ════════════════════════════════════════════════════════════════════════════
# AUTH ROUTES
# ════════════════════════════════════════════════════════════════════════════

@app.route("/signup", methods=["GET", "POST"])
def signup():
    if session.get("user_id"):
        return redirect(url_for("dashboard"))

    if request.method == "POST":
        full_name = request.form.get("full_name", "").strip()
        email     = request.form.get("email", "").strip().lower()
        password  = request.form.get("password", "")
        confirm   = request.form.get("confirm_password", "")
        role      = request.form.get("role", "user")
        github    = request.form.get("github", "").strip()
        linkedin  = request.form.get("linkedin", "").strip()

        errors = []
        if len(full_name) < 2:
            errors.append("Please enter your full name.")
        if not re.match(r"^[^\s@]+@[^\s@]+\.[^\s@]+$", email):
            errors.append("Enter a valid email address.")
        if len(password) < 8:
            errors.append("Password must be at least 8 characters.")
        if password != confirm:
            errors.append("Passwords do not match.")
        if email in USERS:
            errors.append("An account with this email already exists.")
        if role not in ("user", "recruiter"):
            role = "user"

        if errors:
            for e in errors:
                flash(e, "error")
            return render_template("signup.html")

        USERS[email] = {
            "id":            email,
            "full_name":     full_name,
            "email":         email,
            "password_hash": generate_password_hash(password),
            "role":          role,
            "github":        github,
            "linkedin":      linkedin,
            "created_at":    datetime.utcnow(),
        }
        flash("Account created! Welcome aboard.", "success")
        session.clear()
        session["user_id"] = email
        if role == "recruiter":
            return redirect(url_for("recruiter"))
        return redirect(url_for("dashboard"))

    return render_template("signup.html")


@app.route("/login", methods=["GET", "POST"])
def login():
    if session.get("user_id"):
        u = current_user()
        if u and u.get("role") == "recruiter":
            return redirect(url_for("recruiter"))
        return redirect(url_for("dashboard"))

    if request.method == "POST":
        email    = request.form.get("email", "").strip().lower()
        password = request.form.get("password", "")

        user = USERS.get(email)
        if user and check_password_hash(user["password_hash"], password):
            session.clear()
            session["user_id"] = email
            session.permanent = bool(request.form.get("remember"))
            flash(f"Welcome back, {user['full_name'].split()[0]}!", "success")
            if user["role"] == "recruiter":
                return redirect(url_for("recruiter"))
            # Honor ?next= param (only allow internal paths)
            next_url = request.form.get("next") or request.args.get("next", "")
            if next_url and next_url.startswith("/") and not next_url.startswith("//"):
                return redirect(next_url)
            return redirect(url_for("dashboard"))
        else:
            flash("Incorrect email or password.", "error")

    return render_template("login.html")


@app.route("/logout")
def logout():
    session.clear()
    flash("You've been signed out.", "info")
    return redirect(url_for("login"))


# ════════════════════════════════════════════════════════════════════════════
# DASHBOARD & PROFILE
# ════════════════════════════════════════════════════════════════════════════

@app.route("/dashboard")
@login_required
def dashboard():
    u = current_user()
    if not u:
        session.clear()
        return redirect(url_for("login"))
    if u.get("role") == "recruiter":
        return redirect(url_for("recruiter"))
    return render_template("dashboard.html",
        greeting=get_greeting(),
        analyses_count=session.get("analyses_count", 0),
        best_score=session.get("best_score", None),
        interview_sets=0,
        recent_analyses=[],
    )


@app.route("/profile", methods=["GET", "POST"])
@login_required
def profile():
    u = current_user()
    if not u:
        return redirect(url_for("login"))

    if request.method == "POST":
        full_name    = request.form.get("full_name", "").strip()
        email        = request.form.get("email", "").strip().lower()
        github       = request.form.get("github", "").strip()
        linkedin     = request.form.get("linkedin", "").strip()
        new_password = request.form.get("new_password", "")

        errors = []
        if len(full_name) < 2:
            errors.append("Name is too short.")
        if not re.match(r"^[^\s@]+@[^\s@]+\.[^\s@]+$", email):
            errors.append("Invalid email address.")
        if email != u["email"] and email in USERS:
            errors.append("That email is already in use.")

        if errors:
            for e in errors:
                flash(e, "error")
            return render_template("profile.html")

        old_email = u["email"]
        u["full_name"] = full_name
        u["github"]    = github
        u["linkedin"]  = linkedin
        if new_password:
            if len(new_password) < 8:
                flash("New password must be at least 8 characters.", "error")
                return render_template("profile.html")
            u["password_hash"] = generate_password_hash(new_password)
        if email != old_email:
            USERS[email] = u
            del USERS[old_email]
            session["user_id"] = email
            u["email"] = email
        else:
            USERS[old_email] = u

        flash("Profile updated!", "success")
        return redirect(url_for("profile"))

    return render_template("profile.html")


# ════════════════════════════════════════════════════════════════════════════
# PUBLIC ROUTES
# ════════════════════════════════════════════════════════════════════════════

@app.route("/")
def index():
    if session.get("user_id"):
        u = current_user()
        if u is None:
            session.clear()
        elif u.get("role") == "recruiter":
            return redirect(url_for("recruiter"))
        else:
            return redirect(url_for("dashboard"))
    return render_template("index.html")


@app.route("/about")
def about():
    return render_template("about.html")


# ════════════════════════════════════════════════════════════════════════════
# ATS ANALYZER
# ════════════════════════════════════════════════════════════════════════════

@app.route("/analyze", methods=["GET"])
@user_required
def analyze():
    return render_template("analyze.html",
        categories=["None"] + list(JOB_ROLES.keys()))


@app.route("/get_roles")
def get_roles():
    cat = request.args.get("category", "None")
    return jsonify([] if cat == "None" or cat not in JOB_ROLES
                   else list(JOB_ROLES[cat].keys()))


@app.route("/analyze", methods=["POST"])
@user_required
def analyze_post():
    if "resume" not in request.files or request.files["resume"].filename == "":
        flash("Please select a file.", "error")
        return redirect(url_for("analyze"))
    f = request.files["resume"]
    if not allowed_file(f.filename):
        flash("Only PDF/DOCX accepted.", "error")
        return redirect(url_for("analyze"))

    fn = secure_filename(f.filename)
    fp = os.path.join(app.config["UPLOAD_FOLDER"], fn)
    try:
        f.save(fp)
        with open(fp, "rb") as fh:
            class _W:
                def __init__(self, x, n): self._f = x; self.name = n
                def read(self):           return self._f.read()
                def seek(self, *a):       return self._f.seek(*a)
            parsed = parser.parse(_W(fh, fn))
            text   = parsed.get("raw_text", "").strip()
    except Exception as e:
        flash(f"Read error: {e}", "error")
        return redirect(url_for("analyze"))
    finally:
        if os.path.exists(fp): os.remove(fp)

    if not text:
        flash("No text extracted from the file.", "error")
        return redirect(url_for("analyze"))

    cat  = request.form.get("category", "None")
    role = request.form.get("role", "None")
    if cat == "None" or role == "None":
        ri = {"description": "General analysis.", "required_skills": []}
        dr = "General Analysis"; dc = "General"
    else:
        ri = JOB_ROLES.get(cat, {}).get(role, {"description": "", "required_skills": []})
        dr = role; dc = cat

    try:
        analysis = analyzer.analyze_resume({"raw_text": text}, ri)
    except Exception as e:
        flash(f"Analysis error: {e}", "error")
        return redirect(url_for("analyze"))

    ml = get_ats_ml(text)
    if ml is not None:
        analysis["ats_score"] = ml

    if analysis.get("document_type") not in ("resume", None):
        flash(f"Looks like a '{analysis['document_type']}' — upload a resume.", "warning")
        return redirect(url_for("analyze"))

    # Track stats in session
    score = int(analysis.get("ats_score", 0))
    session["analyses_count"] = session.get("analyses_count", 0) + 1
    best = session.get("best_score")
    if best is None or score > best:
        session["best_score"] = score

    return render_template("result.html", analysis=analysis,
                           role=dr, category=dc, role_info=ri)


# ════════════════════════════════════════════════════════════════════════════
# RESUME BUILDER
# ════════════════════════════════════════════════════════════════════════════

@app.route("/builder", methods=["GET"])
@user_required
def builder():
    return render_template("builder.html")


@app.route("/builder", methods=["POST"])
@user_required
def builder_post():
    """Build a DOCX resume from form data and stream it as a download."""
    if not DOCX_LOADED:
        flash("python-docx is not installed. Run: pip install python-docx", "error")
        return redirect(url_for("builder"))

    data      = request.form
    full_name = data.get("full_name", "").strip()
    email_val = data.get("email", "").strip()
    phone     = data.get("phone", "").strip()
    location  = data.get("location", "").strip()
    linkedin  = data.get("linkedin", "").strip()
    portfolio = data.get("portfolio", "").strip()
    summary   = data.get("summary", "").strip()

    if not full_name or not email_val:
        flash("Full name and email are required.", "error")
        return redirect(url_for("builder"))

    def _parse_json(key, default):
        try:
            return json.loads(data.get(key, default))
        except Exception:
            return json.loads(default)

    experiences = _parse_json("experiences_json", "[]")
    educations  = _parse_json("education_json",   "[]")
    projects    = _parse_json("projects_json",    "[]")
    skills_data = _parse_json("skills_json",      "{}")

    # ── Build DOCX ────────────────────────────────────────────────────────
    doc = DocxDocument()

    for sec in doc.sections:
        sec.top_margin    = Inches(0.75)
        sec.bottom_margin = Inches(0.75)
        sec.left_margin   = Inches(1.0)
        sec.right_margin  = Inches(1.0)

    def _heading(text):
        p   = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text.upper())
        run.bold = True
        run.font.size = Pt(11)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after  = Pt(2)
        pPr  = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bot  = OxmlElement("w:bottom")
        bot.set(qn("w:val"),   "single")
        bot.set(qn("w:sz"),    "6")
        bot.set(qn("w:space"), "1")
        bot.set(qn("w:color"), "2196F3")
        pBdr.append(bot)
        pPr.append(pBdr)

    def _bullet(text):
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(text).font.size = Pt(10)
        p.paragraph_format.space_after = Pt(1)

    def _body(text):
        p = doc.add_paragraph(text)
        if p.runs:
            p.runs[0].font.size = Pt(10)
        p.paragraph_format.space_after = Pt(2)

    # Name
    np_ = doc.add_paragraph()
    np_.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nr  = np_.add_run(full_name)
    nr.bold = True; nr.font.size = Pt(20)

    # Contact line
    contact = " | ".join(x for x in [email_val, phone, location, linkedin, portfolio] if x)
    cp = doc.add_paragraph(contact)
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if cp.runs:
        cp.runs[0].font.size = Pt(9)
    cp.paragraph_format.space_after = Pt(4)

    # Summary
    if summary:
        _heading("Professional Summary")
        _body(summary)

    # Experience
    if experiences:
        _heading("Work Experience")
        for exp in experiences:
            p = doc.add_paragraph()
            p.add_run(exp.get("position", "")).bold = True
            p.add_run(f"  —  {exp.get('company', '')}")
            dr = f"{exp.get('start_date', '')} – {exp.get('end_date', 'Present')}"
            tr = p.add_run(f"\t{dr}"); tr.font.size = Pt(9)
            p.paragraph_format.space_after = Pt(1)
            if exp.get("description"): _body(exp["description"])
            for r in exp.get("responsibilities", []):
                if r: _bullet(r)
            for a in exp.get("achievements", []):
                if a: _bullet(f"✓ {a}")

    # Education
    if educations:
        _heading("Education")
        for edu in educations:
            p = doc.add_paragraph()
            p.add_run(f"{edu.get('degree', '')} in {edu.get('field', '')}").bold = True
            p.add_run(f"  —  {edu.get('school', '')}")
            meta = "  |  ".join(x for x in [
                edu.get("graduation_date", ""),
                f"GPA: {edu.get('gpa')}" if edu.get("gpa") else ""
            ] if x)
            if meta:
                tr = p.add_run(f"\t{meta}"); tr.font.size = Pt(9)
            for a in edu.get("achievements", []):
                if a: _bullet(a)

    # Projects
    if projects:
        _heading("Projects")
        for proj in projects:
            p = doc.add_paragraph()
            p.add_run(proj.get("name", "")).bold = True
            if proj.get("technologies"):
                tr = p.add_run(f"  |  {proj['technologies']}"); tr.font.size = Pt(9)
            if proj.get("link"):
                tr = p.add_run(f"  [{proj['link']}]"); tr.font.size = Pt(9)
            if proj.get("description"): _body(proj["description"])
            for r in proj.get("responsibilities", []):
                if r: _bullet(r)
            for a in proj.get("achievements", []):
                if a: _bullet(f"✓ {a}")

    # Skills
    if skills_data:
        _heading("Skills")
        for label, key in [("Technical", "technical"), ("Soft Skills", "soft"),
                            ("Languages", "languages"), ("Tools & Tech", "tools")]:
            vals = skills_data.get(key, [])
            if vals:
                p = doc.add_paragraph()
                p.add_run(f"{label}: ").bold = True
                p.add_run(", ".join(vals)).font.size = Pt(10)
                p.paragraph_format.space_after = Pt(1)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    filename = f"{full_name.replace(' ', '_')}_resume.docx"
    return send_file(
        buf,
        as_attachment=True,
        download_name=filename,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


# ════════════════════════════════════════════════════════════════════════════
# AI INTERVIEW QUESTIONS
# ════════════════════════════════════════════════════════════════════════════

ENCOURAGING_MSGS = [
    "All the best for your interview!",
    "Hope these help you prepare well!",
    "Good luck with your interview preparation!",
    "You're going to ace this!",
    "Keep up the great work!",
    "Believe in yourself and your abilities!",
    "Remember, preparation is key to success!",
    "You've got this! Go show them what you're made of!",
    "Stay confident and positive!",
    "Every question is a step towards your goal!",
    "Practice makes perfect!",
    "You're one step closer to your dream job!",
    "Keep pushing forward!",
    "Success is just around the corner!",
    "Your hard work will pay off!",
]


@app.route("/ai-questions", methods=["GET", "POST"])
@user_required
def ai_questions():
    questions   = []
    prompt      = ""
    skills_str  = ""
    error       = None
    warning     = None
    encouraging = None

    if request.method == "POST":
        prompt = request.form.get("prompt", "").strip()
        try:
            from utils.ai_engine import get_questions_from_prompt
            result     = get_questions_from_prompt(prompt)
            questions  = result.get("questions", [])
            skills_str = result.get("skills_str", "")
            warning    = result.get("warning")
            error      = result.get("error")
            if questions:
                encouraging = random.choice(ENCOURAGING_MSGS)
        except Exception as exc:
            logger.error(f"ai_questions error: {exc}")
            error = f"Question generation failed: {exc}"

    return render_template("ai_questions.html",
        prompt=prompt, questions=questions, skills_str=skills_str,
        error=error, warning=warning, encouraging=encouraging)


# ════════════════════════════════════════════════════════════════════════════
# JOB SEARCH
# ════════════════════════════════════════════════════════════════════════════

@app.route("/job-search", methods=["GET", "POST"])
@user_required
def job_search():
    results    = []
    query      = ""
    location_q = ""
    error      = None

    if request.method == "POST":
        query      = request.form.get("query", "").strip()
        location_q = request.form.get("location", "").strip()

        if not query:
            error = "Please enter a job title or keyword."
        else:
            app_id  = os.environ.get("ADZUNA_APP_ID",  "")
            app_key = os.environ.get("ADZUNA_APP_KEY", "")

            if app_id and app_key:
                try:
                    import requests as req_lib
                    from urllib.parse import quote as url_quote
                    country = "in"   # change to 'gb', 'us', etc. as needed
                    api_url = (
                        f"https://api.adzuna.com/v1/api/jobs/{country}/search/1"
                        f"?app_id={app_id}&app_key={app_key}"
                        f"&results_per_page=20"
                        f"&what={url_quote(query)}"
                        f"&where={url_quote(location_q)}"
                        f"&content-type=application/json"
                    )
                    resp = req_lib.get(api_url, timeout=8)
                    if resp.ok:
                        for job in resp.json().get("results", []):
                            results.append({
                                "title":    job.get("title", ""),
                                "company":  job.get("company",  {}).get("display_name", "N/A"),
                                "location": job.get("location", {}).get("display_name", "N/A"),
                                "salary":   _fmt_salary(job),
                                "url":      job.get("redirect_url", "#"),
                                "desc":     (job.get("description", "")[:200] + "…"),
                                "date":     job.get("created", "")[:10],
                            })
                    else:
                        error = f"Job API error ({resp.status_code}). Please try again later."
                except Exception as exc:
                    error = f"Could not fetch jobs: {exc}"
            else:
                error = (
                    "Job search requires Adzuna API credentials. "
                    "Set ADZUNA_APP_ID and ADZUNA_APP_KEY environment variables, "
                    "or visit adzuna.com to get free API keys."
                )

    return render_template("job_search.html",
        results=results, query=query, location_q=location_q, error=error)


# ════════════════════════════════════════════════════════════════════════════
# FEEDBACK
# ════════════════════════════════════════════════════════════════════════════

@app.route("/feedback", methods=["GET"])
@login_required
def feedback():
    avg_rating = 0
    if _feedback_store:
        avg_rating = round(
            sum(f["rating"] for f in _feedback_store) / len(_feedback_store), 1
        )
    dist = {1: 0, 2: 0, 3: 0, 4: 0, 5: 0}
    for f in _feedback_store:
        dist[f["rating"]] = dist.get(f["rating"], 0) + 1

    return render_template("feedback.html",
        feedbacks=_feedback_store,
        avg_rating=avg_rating,
        rating_dist=dist,
        total=len(_feedback_store),
    )


@app.route("/feedback", methods=["POST"])
@login_required
def feedback_post():
    u       = current_user()
    name    = request.form.get("name",           u.get("full_name", "")).strip()
    email_v = request.form.get("feedback_email", u.get("email", "")).strip()
    rating  = max(1, min(5, int(request.form.get("rating", 3))))
    feature = request.form.get("feature", "").strip()
    comment = request.form.get("comment", "").strip()

    if not comment:
        flash("Please write a comment before submitting.", "error")
        return redirect(url_for("feedback"))

    _feedback_store.append({
        "name":    name,
        "email":   email_v,
        "rating":  rating,
        "feature": feature,
        "comment": comment,
        "date":    datetime.utcnow().strftime("%d %b %Y"),
    })
    flash("Thank you for your feedback! 🙏", "success")
    return redirect(url_for("feedback"))


# ════════════════════════════════════════════════════════════════════════════
# RECRUITER ROUTES
# ════════════════════════════════════════════════════════════════════════════

# ── Recruiter keyword constants ──────────────────────────────────────────────
EDU_KW = {
    "Schooling":      ["10th","high school","ssc","school","secondary","matric"],
    "Intermediate":   ["12th","intermediate","hsc","senior secondary","puc"],
    "Bachelors":      ["btech","b.tech","bsc","be","b.e","bachelor","undergraduate","degree","b.a","b.com"],
    "Masters":        ["mtech","m.tech","msc","mba","master","postgraduate","m.s","m.a"],
    "PhD":            ["phd","ph.d","doctorate"],
    "Certifications": ["diploma","certificate","certification"],
}
TECH_SKILLS = {
    "Prog":   ["c","c++","java","python","javascript","ruby","go","kotlin","swift","php","scala"],
    "Web":    ["html","css","node.js","react","django","flask","angular","vue.js","bootstrap"],
    "DB":     ["mysql","mongodb","postgresql","sqlite","oracle","nosql"],
    "FW":     ["pandas","numpy","tensorflow","pytorch","scikit-learn","spring"],
    "Cloud":  ["aws","google cloud","azure","docker","kubernetes"],
    "Tools":  ["git","github","jenkins","jira","linux","android studio"],
    "Mobile": ["android","ios","flutter","react native"],
}
SCORE_PAT = re.compile(r"(?i)(?:gpa|cgpa|percentage|marks)[:\s-]*([\d\.]+(?:/[\d\.]+)?%?)")


def _extract_text_r(fbytes, ftype):
    try:
        if "pdf" in ftype and PDFMINER_LOADED:
            return pdfminer_extract(io.BytesIO(fbytes))
        if "word" in ftype and DOCX_LOADED:
            doc = DocxDocument(io.BytesIO(fbytes))
            return "\n".join(p.text for p in doc.paragraphs)
        import PyPDF2
        r = PyPDF2.PdfReader(io.BytesIO(fbytes))
        return "\n".join(p.extract_text() or "" for p in r.pages)
    except Exception:
        return ""


def _is_resume(text):
    tl  = text.lower()
    scs = sum(s in tl for s in ["education","experience","skills","projects","internships"])
    kws = [k for v in EDU_KW.values() for k in v] + [s for v in TECH_SKILLS.values() for s in v]
    return scs >= 2 or sum(k in tl for k in kws) >= 5


def _norm_score(s_str, _level):
    try:
        s_str = s_str.lower()
        if "%" in s_str:
            m = re.search(r"(\d+\.?\d*)%", s_str); return float(m.group(1)) if m else None
        m = re.search(r"(\d+\.?\d*)\s*/\s*(\d+\.?\d*)", s_str)
        if m:
            s, mx = float(m.group(1)), float(m.group(2)); return (s / mx) * 100
        m = re.search(r"(\d+\.?\d*)", s_str)
        if m:
            s = float(m.group(1))
            if s <= 4:  return (s / 4)  * 100
            if s <= 10: return (s / 10) * 100
            return s if s <= 100 else None
    except Exception:
        return None


def _extract_edu(text):
    details = []
    for line in text.split("\n"):
        ll  = line.strip().lower()
        if not ll: continue
        deg = next((cat for cat, kws in EDU_KW.items() if any(k in ll for k in kws)), "")
        if not deg: continue
        sm  = SCORE_PAT.search(line)
        sr  = sm.group(0) if sm else "N/A"
        ns  = _norm_score(sr, deg) if sm else None
        ym  = re.search(r"\b(?:19|20)\d{2}\b", line)
        details.append({"degree": deg, "score": sr, "normalized_score": ns,
                        "year": ym.group(0) if ym else ""})
    return details


def _extract_skills_r(text, required):
    skills = set(); tl = text.lower()
    for sl in TECH_SKILLS.values():
        for s in sl:
            if re.search(rf"\b{re.escape(s)}\b", tl): skills.add(s)
    for s in required:
        ss = s if isinstance(s, str) else " or ".join(s)
        if re.search(rf"\b{re.escape(ss)}\b", tl):
            skills.update(s if isinstance(s, list) else [s])
    return skills


def _parse_criteria(q):
    edu, skills, thresh = {}, [], 0
    pat = re.compile(
        r"(schooling|intermediate|bachelors?|masters?|phd|certifications?)\s*"
        r"(?:above|with)?\s*(\d+\.?\d*)\s*(%|gpa)?", re.I)
    for m in pat.finditer(q):
        lv, sc, un = m.group(1).capitalize(), float(m.group(2)), (m.group(3) or "").lower()
        pct = (sc if un == "%" else
               ((sc / 4) * 100 if un == "gpa" and sc <= 4 else
                (sc / 10) * 100 if un == "gpa" else sc))
        edu[lv] = max(0, min(pct, 100))
    pool = {s for v in TECH_SKILLS.values() for s in v}
    for w in q.lower().split():
        if w in pool: skills.append(w)
    am = re.search(r"ats\s*(?:score)?\s*(?:above)?\s*(\d+)", q, re.I)
    if am: thresh = int(am.group(1))
    return {"edu": edu, "skills": list(set(skills)), "thresh": thresh}


def _score_candidate(name, text, crit):
    edu    = _extract_edu(text)
    skills = _extract_skills_r(text, crit["skills"])
    ml     = get_ats_ml(text)
    if ml is None:
        r  = analyzer.analyze_resume({"raw_text": text}, {"required_skills": crit["skills"]})
        ml = r.get("ats_score", 50)
    ats    = round(float(ml), 1)
    m_edu  = (not crit["edu"]) or all(
        any(e["degree"].lower() == lv.lower() and
            e["normalized_score"] and e["normalized_score"] >= req for e in edu)
        for lv, req in crit["edu"].items())
    m_sk   = (not crit["skills"]) or all(
        (any(s in [x.lower() for x in skills] for s in sk) if isinstance(sk, list)
         else sk.lower() in [s.lower() for s in skills])
        for sk in crit["skills"])
    m_ats  = ats >= crit["thresh"] if crit["thresh"] else True
    em     = (sum(
        any(e["degree"].lower() == lv.lower() and
            e["normalized_score"] and e["normalized_score"] >= req for e in edu)
        for lv, req in crit["edu"].items()) / len(crit["edu"])) if crit["edu"] else 1.0
    sm_val = (sum(
        (any(s in [x.lower() for x in skills] for s in sk)
         if isinstance(sk, list) else sk.lower() in [s.lower() for s in skills])
        for sk in crit["skills"]) / len(crit["skills"])) if crit["skills"] else 1.0
    overall = round((em + sm_val + (1 if m_ats else 0)) / 3 * 100, 1)
    return {
        "name":    name,
        "ats":     ats,
        "match":   overall,
        "edu":     "; ".join(f"{e['degree']} ({e['score']})" for e in edu) or "N/A",
        "skills":  ", ".join(sorted(skills)) or "N/A",
        "edu_met": m_edu,
        "sk_met":  m_sk,
        "ats_met": m_ats,
        "ok":      (m_edu and m_sk and m_ats) or overall >= 50,
    }


@app.route("/recruiter")
@recruiter_required
def recruiter():
    return render_template("recruiter.html", candidates=session.get("candidates", []))


@app.route("/recruiter/screen", methods=["POST"])
@recruiter_required
def rec_screen():
    files = request.files.getlist("resumes")
    query = request.form.get("criteria", "").strip()
    if not files or all(f.filename == "" for f in files):
        flash("Upload at least one resume.", "error")
        return redirect(url_for("recruiter"))
    crit = _parse_criteria(query) if query else {"edu": {}, "skills": [], "thresh": 0}
    results = []; skipped = []

    def _proc(uf):
        nm   = secure_filename(uf.filename)
        fb   = uf.read()
        ft   = uf.content_type or ""
        text = _extract_text_r(fb, ft)
        if not text or not _is_resume(text): return None, nm
        return _score_candidate(nm, text, crit), None

    with ThreadPoolExecutor(max_workers=4) as ex:
        for r, s in ex.map(_proc, [f for f in files if f.filename != ""]):
            if r: results.append(r)
            if s: skipped.append(s)

    results.sort(key=lambda x: x["match"], reverse=True)
    session["candidates"] = results
    if skipped: flash(f"Skipped (not resumes): {', '.join(skipped)}", "warning")
    flash(f"Processed {len(results)} resume(s) successfully.", "success")
    return redirect(url_for("recruiter"))


@app.route("/recruiter/export")
@recruiter_required
def rec_export():
    import csv
    cands = session.get("candidates", [])
    if not cands:
        flash("Nothing to export.", "warning")
        return redirect(url_for("recruiter"))
    si = io.StringIO()
    w  = csv.DictWriter(si, fieldnames=cands[0].keys())
    w.writeheader(); w.writerows(cands)
    return Response(si.getvalue(), mimetype="text/csv",
                    headers={"Content-Disposition": "attachment;filename=shortlisted.csv"})


# ════════════════════════════════════════════════════════════════════════════
# ERROR HANDLERS
# ════════════════════════════════════════════════════════════════════════════

@app.errorhandler(404)
def e404(e):
    return render_template("error.html", code=404, message="Page not found."), 404

@app.errorhandler(413)
def e413(e):
    flash("File too large (max 10 MB).", "error")
    return redirect(url_for("index"))

@app.errorhandler(500)
def e500(e):
    return render_template("error.html", code=500, message="Something went wrong on our end."), 500


# ════════════════════════════════════════════════════════════════════════════
# AI ENHANCEMENT ROUTES  –  added by Smart Resume Analyzer AI v2
# ════════════════════════════════════════════════════════════════════════════

from utils.ai_engine import (
    get_smart_feedback,
    get_skill_gaps,
    get_interview_questions,
    get_rewrite_suggestions,
    get_role_optimization,
)


# ── AI Feedback ──────────────────────────────────────────────────────────────

@app.route("/ai-feedback", methods=["GET", "POST"])
@user_required
def ai_feedback():
    result = None
    error = None

    if request.method == "POST":
        if "resume" not in request.files or request.files["resume"].filename == "":
            flash("Please upload a resume.", "error")
            return redirect(url_for("ai_feedback"))

        f = request.files["resume"]
        if not allowed_file(f.filename):
            flash("Only PDF/DOCX accepted.", "error")
            return redirect(url_for("ai_feedback"))

        fn = secure_filename(f.filename)
        fp = os.path.join(app.config["UPLOAD_FOLDER"], fn)
        try:
            f.save(fp)
            with open(fp, "rb") as fh:
                class _W:
                    def __init__(self, x, n): self._f = x; self.name = n
                    def read(self):           return self._f.read()
                    def seek(self, *a):       return self._f.seek(*a)
                parsed = parser.parse(_W(fh, fn))
                text = parsed.get("raw_text", "").strip()
        except Exception as e:
            error = f"Could not read file: {e}"
            return render_template("ai_feedback.html", result=None, error=error,
                                   categories=list(JOB_ROLES.keys()))
        finally:
            if os.path.exists(fp): os.remove(fp)

        if not text:
            error = "No text could be extracted from the file."
            return render_template("ai_feedback.html", result=None, error=error,
                                   categories=list(JOB_ROLES.keys()))

        cat  = request.form.get("category", "None")
        role = request.form.get("role", "None")
        ri   = {} if cat == "None" or role == "None" else JOB_ROLES.get(cat, {}).get(role, {})
        required = ri.get("required_skills", [])
        job_role = role if role != "None" else ""

        try:
            result = get_smart_feedback(text, job_role, required)
            result["resume_text"] = text[:500]  # preview only
        except Exception as e:
            logger.error(f"AI feedback error: {e}")
            error = f"Analysis error: {e}"

    return render_template("ai_feedback.html", result=result, error=error,
                           categories=list(JOB_ROLES.keys()))


# ── Skill Gap Analysis ───────────────────────────────────────────────────────

@app.route("/skill-gap", methods=["GET", "POST"])
@user_required
def skill_gap():
    result = None
    error = None

    if request.method == "POST":
        if "resume" not in request.files or request.files["resume"].filename == "":
            flash("Please upload a resume.", "error")
            return redirect(url_for("skill_gap"))

        f = request.files["resume"]
        if not allowed_file(f.filename):
            flash("Only PDF/DOCX accepted.", "error")
            return redirect(url_for("skill_gap"))

        fn = secure_filename(f.filename)
        fp = os.path.join(app.config["UPLOAD_FOLDER"], fn)
        try:
            f.save(fp)
            with open(fp, "rb") as fh:
                class _W:
                    def __init__(self, x, n): self._f = x; self.name = n
                    def read(self):           return self._f.read()
                    def seek(self, *a):       return self._f.seek(*a)
                parsed = parser.parse(_W(fh, fn))
                text = parsed.get("raw_text", "").strip()
        except Exception as e:
            error = f"Could not read file: {e}"
            return render_template("skill_gap.html", result=None, error=error,
                                   categories=list(JOB_ROLES.keys()))
        finally:
            if os.path.exists(fp): os.remove(fp)

        cat  = request.form.get("category", "None")
        role = request.form.get("role", "None")
        ri   = {} if cat == "None" or role == "None" else JOB_ROLES.get(cat, {}).get(role, {})
        required = ri.get("required_skills", [])
        job_role = role if role != "None" else "Software Engineer"

        try:
            result = get_skill_gaps(text, job_role, required, cat)
        except Exception as e:
            logger.error(f"Skill gap error: {e}")
            error = f"Analysis error: {e}"

    return render_template("skill_gap.html", result=result, error=error,
                           categories=list(JOB_ROLES.keys()))


# ── Smart Interview Questions ────────────────────────────────────────────────

@app.route("/smart-questions", methods=["GET", "POST"])
@user_required
def smart_questions():
    result = None
    error = None

    if request.method == "POST":
        cat   = request.form.get("category", "None")
        role  = request.form.get("role", "None")
        level = request.form.get("level", "mid")
        num   = min(int(request.form.get("num_questions", 15)), 30)

        ri = {} if cat == "None" or role == "None" else JOB_ROLES.get(cat, {}).get(role, {})
        skills = ri.get("required_skills", [])
        job_role = role if role != "None" else "Software Engineer"

        try:
            result = get_interview_questions(job_role, skills, level, num)
        except Exception as e:
            logger.error(f"Smart questions error: {e}")
            error = f"Generation error: {e}"

    return render_template("smart_questions.html", result=result, error=error,
                           categories=list(JOB_ROLES.keys()))


# ── Resume Rewrite ───────────────────────────────────────────────────────────

@app.route("/rewrite", methods=["GET", "POST"])
@user_required
def rewrite():
    result = None
    error = None

    if request.method == "POST":
        if "resume" not in request.files or request.files["resume"].filename == "":
            flash("Please upload a resume.", "error")
            return redirect(url_for("rewrite"))

        f = request.files["resume"]
        if not allowed_file(f.filename):
            flash("Only PDF/DOCX accepted.", "error")
            return redirect(url_for("rewrite"))

        fn = secure_filename(f.filename)
        fp = os.path.join(app.config["UPLOAD_FOLDER"], fn)
        try:
            f.save(fp)
            with open(fp, "rb") as fh:
                class _W:
                    def __init__(self, x, n): self._f = x; self.name = n
                    def read(self):           return self._f.read()
                    def seek(self, *a):       return self._f.seek(*a)
                parsed = parser.parse(_W(fh, fn))
                text = parsed.get("raw_text", "").strip()
        except Exception as e:
            error = f"Could not read file: {e}"
            return render_template("rewrite.html", result=None, error=error,
                                   categories=list(JOB_ROLES.keys()))
        finally:
            if os.path.exists(fp): os.remove(fp)

        cat  = request.form.get("category", "None")
        role = request.form.get("role", "None")
        job_role = role if role != "None" else ""

        try:
            result = get_rewrite_suggestions(text, job_role)
        except Exception as e:
            logger.error(f"Rewrite error: {e}")
            error = f"Analysis error: {e}"

    return render_template("rewrite.html", result=result, error=error,
                           categories=list(JOB_ROLES.keys()))


# ── Role Optimizer ───────────────────────────────────────────────────────────

@app.route("/role-optimizer", methods=["GET", "POST"])
@user_required
def role_optimizer():
    result = None
    error = None

    if request.method == "POST":
        if "resume" not in request.files or request.files["resume"].filename == "":
            flash("Please upload a resume.", "error")
            return redirect(url_for("role_optimizer"))

        f = request.files["resume"]
        if not allowed_file(f.filename):
            flash("Only PDF/DOCX accepted.", "error")
            return redirect(url_for("role_optimizer"))

        fn = secure_filename(f.filename)
        fp = os.path.join(app.config["UPLOAD_FOLDER"], fn)
        try:
            f.save(fp)
            with open(fp, "rb") as fh:
                class _W:
                    def __init__(self, x, n): self._f = x; self.name = n
                    def read(self):           return self._f.read()
                    def seek(self, *a):       return self._f.seek(*a)
                parsed = parser.parse(_W(fh, fn))
                text = parsed.get("raw_text", "").strip()
        except Exception as e:
            error = f"Could not read file: {e}"
            return render_template("role_optimizer.html", result=None, error=error,
                                   categories=list(JOB_ROLES.keys()))
        finally:
            if os.path.exists(fp): os.remove(fp)

        cat  = request.form.get("category", "None")
        role = request.form.get("role", "None")
        ri   = {} if cat == "None" or role == "None" else JOB_ROLES.get(cat, {}).get(role, {})
        required = ri.get("required_skills", [])
        job_role = role if role != "None" else "Software Engineer"

        try:
            result = get_role_optimization(text, job_role, required, cat)
        except Exception as e:
            logger.error(f"Role optimizer error: {e}")
            error = f"Analysis error: {e}"

    return render_template("role_optimizer.html", result=result, error=error,
                           categories=list(JOB_ROLES.keys()))


# ── API endpoint: get roles (also used by new pages) ─────────────────────────
# (already defined in original app.py as /get_roles)

# ════════════════════════════════════════════════════════════════════════════
if __name__ == "__main__":
    app.run(debug=True, host="0.0.0.0", port=5000)
