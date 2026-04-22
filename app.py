import os, io, re, json, math, random, logging, joblib, tempfile, secrets

# Load .env file if present
try:
    from dotenv import load_dotenv
    load_dotenv()
except ImportError:
    pass

from datetime import datetime, timedelta
from functools import wraps
from concurrent.futures import ThreadPoolExecutor

from flask import (Flask, render_template, request, redirect,
                   url_for, session, flash, jsonify, Response, send_file)
from werkzeug.utils import secure_filename
from werkzeug.security import generate_password_hash, check_password_hash

from utils.resume_parser import ResumeParser
from utils.resume_analyzer import ResumeAnalyzer
from config.job_roles import JOB_ROLES

# ── Optional dependencies ────────────────────────────────────────────────────
try:
    import spacy; nlp = spacy.load("en_core_web_sm"); SPACY_LOADED = True
except Exception:
    nlp = None; SPACY_LOADED = False
try:
    from pdfminer.high_level import extract_text as pdfminer_extract; PDFMINER_LOADED = True
except Exception:
    PDFMINER_LOADED = False
try:
    from docx import Document as DocxDocument
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_LOADED = True
except Exception:
    DOCX_LOADED = False
try:
    import pandas as pd; PANDAS_LOADED = True
except Exception:
    PANDAS_LOADED = False

# ── App setup ────────────────────────────────────────────────────────────────
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

app = Flask(__name__)

# ── Secret key — robust for production ──────────────────────────────────────
# On HF Spaces: set FLASK_SECRET_KEY in Space secrets.
# Fallback: derive from a file so it survives container restarts within the
# same deployment (but new deployments will invalidate old sessions — fine).
_SECRET_KEY_ENV = os.environ.get("FLASK_SECRET_KEY", "")
if _SECRET_KEY_ENV:
    app.secret_key = _SECRET_KEY_ENV
else:
    _secret_file = os.path.join(
        os.environ.get("DATA_DIR", os.path.join(BASE_DIR, "data")),
        ".secret_key"
    )
    try:
        os.makedirs(os.path.dirname(_secret_file), exist_ok=True)
        if os.path.exists(_secret_file):
            with open(_secret_file) as _sf:
                app.secret_key = _sf.read().strip()
        else:
            _key = secrets.token_hex(32)
            with open(_secret_file, "w") as _sf:
                _sf.write(_key)
            app.secret_key = _key
    except Exception:
        app.secret_key = secrets.token_hex(32)  # volatile fallback
        logger.warning("Could not persist secret key — sessions will reset on restart.")

app.config["PERMANENT_SESSION_LIFETIME"] = timedelta(days=30)
# Make ALL sessions permanent by default — without this, the session cookie
# is a browser-session cookie that HF Spaces' reverse-proxy drops on redirect.
app.config["SESSION_PERMANENT"] = True

# ── Session cookie settings for Hugging Face Spaces (runs behind HTTPS proxy) ─
app.config["SESSION_COOKIE_HTTPONLY"] = True
app.config["SESSION_COOKIE_SAMESITE"] = "None"   # required behind HF proxy
app.config["SESSION_COOKIE_SECURE"]   = True       # required when SameSite=None

# Tell Werkzeug to trust the proxy's X-Forwarded-Proto header so Flask
# knows the request is HTTPS (needed for secure cookies to be sent back).
from werkzeug.middleware.proxy_fix import ProxyFix
app.wsgi_app = ProxyFix(app.wsgi_app, x_proto=1, x_host=1)

# ── Upload folder — use /tmp on read-only filesystems (HF Spaces) ────────────
_UPLOAD_BASE = os.environ.get("UPLOAD_DIR", os.path.join(BASE_DIR, "uploads"))
try:
    os.makedirs(_UPLOAD_BASE, exist_ok=True)
    # Quick write test
    _test = os.path.join(_UPLOAD_BASE, ".write_test")
    with open(_test, "w") as _tf:
        _tf.write("ok")
    os.remove(_test)
    UPLOAD_FOLDER = _UPLOAD_BASE
except Exception:
    UPLOAD_FOLDER = tempfile.gettempdir()
    logger.info(f"Upload folder not writable; falling back to {UPLOAD_FOLDER}")

app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER
app.config["MAX_CONTENT_LENGTH"] = 200 * 1024 * 1024   # 200 MB
USER_MAX_UPLOAD = 10 * 1024 * 1024

# ── ML Models ────────────────────────────────────────────────────────────────
try:
    ats_model  = joblib.load(os.path.join(BASE_DIR, "ats_model.pkl"))
    vectorizer = joblib.load(os.path.join(BASE_DIR, "tfidf_vectorizer.pkl"))
    MODELS_LOADED = True
except Exception as e:
    ats_model = vectorizer = None; MODELS_LOADED = False
    logger.warning(f"ATS model not loaded: {e}")

parser   = ResumeParser()
analyzer = ResumeAnalyzer()

# ── Persistent stores ────────────────────────────────────────────────────────
# On HF Spaces use the /data persistent volume; fall back to local data/
_DATA_DIR      = os.environ.get("DATA_DIR", os.path.join(BASE_DIR, "data"))
_USERS_FILE    = os.path.join(_DATA_DIR, "users.json")
_FEEDBACK_FILE = os.path.join(_DATA_DIR, "feedback.json")
_STATS_FILE    = os.path.join(_DATA_DIR, "stats.json")

try:
    os.makedirs(_DATA_DIR, exist_ok=True)
except Exception as e:
    logger.warning(f"Could not create data dir {_DATA_DIR}: {e}")


def _load_stats():
    try:
        if os.path.exists(_STATS_FILE):
            with open(_STATS_FILE, "r", encoding="utf-8") as fh:
                return json.load(fh)
    except Exception:
        pass
    return {"total_analyses": 0, "best_ats": 0}


def _increment_stats(score: int):
    try:
        stats = _load_stats()
        stats["total_analyses"] = stats.get("total_analyses", 0) + 1
        if score > stats.get("best_ats", 0):
            stats["best_ats"] = score
        with open(_STATS_FILE, "w", encoding="utf-8") as fh:
            json.dump(stats, fh, indent=2)
    except Exception as exc:
        logger.warning(f"Could not update stats: {exc}")


USERS           = {}
_feedback_store = []
import threading
_users_lock = threading.Lock()


def _load_users():
    global USERS
    if os.path.exists(_USERS_FILE):
        try:
            with open(_USERS_FILE, "r", encoding="utf-8") as fh:
                raw = json.load(fh)
            for u in raw.values():
                if isinstance(u.get("created_at"), str):
                    try:
                        u["created_at"] = datetime.fromisoformat(u["created_at"])
                    except Exception:
                        u["created_at"] = datetime.utcnow()
            USERS = raw
        except Exception as e:
            logger.warning(f"Could not load users file: {e}")
            USERS = {}


def _save_users():
    try:
        os.makedirs(_DATA_DIR, exist_ok=True)
        serializable = {}
        for k, v in USERS.items():
            entry = dict(v)
            if isinstance(entry.get("created_at"), datetime):
                entry["created_at"] = entry["created_at"].isoformat()
            serializable[k] = entry
        tmp_file = _USERS_FILE + ".tmp"
        with open(tmp_file, "w", encoding="utf-8") as fh:
            json.dump(serializable, fh, indent=2)
        os.replace(tmp_file, _USERS_FILE)
        logger.info(f"Users saved ({len(serializable)} accounts)")
    except Exception as e:
        logger.error(f"Could not save users file: {e}")


def _load_feedback():
    global _feedback_store
    if os.path.exists(_FEEDBACK_FILE):
        try:
            with open(_FEEDBACK_FILE, "r", encoding="utf-8") as fh:
                _feedback_store = json.load(fh)
        except Exception:
            _feedback_store = []


def _save_feedback():
    try:
        with open(_FEEDBACK_FILE, "w", encoding="utf-8") as fh:
            json.dump(_feedback_store, fh, indent=2)
    except Exception as e:
        logger.warning(f"Could not save feedback: {e}")


_load_users()
_load_feedback()


def _seed_demo():
    changed = False
    for email, name, role in [
        ("user@demo.com",      "Alex Johnson",   "user"),
        ("recruiter@demo.com", "Sarah Mitchell", "recruiter"),
    ]:
        if email not in USERS:
            USERS[email] = {
                "id":            email,
                "full_name":     name,
                "email":         email,
                "password_hash": generate_password_hash("password"),
                "role":          role,
                "github":        "",
                "linkedin":      "",
                "created_at":    datetime.utcnow(),
            }
            changed = True
    if changed:
        _save_users()

_seed_demo()

# ════════════════════════════════════════════════════════════════════════════
# HELPERS & SHARED UTILITIES
# ════════════════════════════════════════════════════════════════════════════

ALLOWED_EXT           = {"pdf", "docx"}
RECRUITER_ALLOWED_EXT = {"pdf", "docx", "csv"}

def allowed_file(fn):
    return "." in fn and fn.rsplit(".", 1)[1].lower() in ALLOWED_EXT

def allowed_recruiter_file(fn):
    return "." in fn and fn.rsplit(".", 1)[1].lower() in RECRUITER_ALLOWED_EXT

def score_color(s):
    s = float(s)
    return "success" if s >= 80 else ("warning" if s >= 60 else "danger")

def get_ats_ml(text):
    if MODELS_LOADED:
        try:
            return round(float(ats_model.predict(vectorizer.transform([text]))[0]), 1)
        except Exception:
            pass
    return None

def get_greeting():
    h = datetime.now().hour
    if h < 12: return "morning"
    if h < 17: return "afternoon"
    return "evening"

def current_user():
    uid = session.get("user_id")
    if not uid:
        return None
    # Always reload from disk so we get the latest user data
    # (handles cross-worker scenarios on HF Spaces)
    with _users_lock:
        _load_users()
    return USERS.get(uid)

def _fmt_salary(job):
    lo = job.get("salary_min")
    hi = job.get("salary_max")
    if lo and hi:
        return f"₹{int(lo):,} – ₹{int(hi):,}"
    if lo:
        return f"from ₹{int(lo):,}"
    return "Not specified"

# ── File save helper — always works even on read-only base paths ──────────────
def _save_upload(file_storage) -> str:
    """Save an uploaded file to a writable temp location; return the path."""
    fn  = secure_filename(file_storage.filename)
    # Try configured upload folder first; fall back to system tmp
    for folder in [app.config["UPLOAD_FOLDER"], tempfile.gettempdir()]:
        try:
            os.makedirs(folder, exist_ok=True)
            # Use a unique name to avoid collisions on multi-worker deployments
            unique_fn = f"{secrets.token_hex(8)}_{fn}"
            fp = os.path.join(folder, unique_fn)
            file_storage.save(fp)
            return fp
        except Exception as e:
            logger.warning(f"Could not save to {folder}: {e}")
    raise RuntimeError("No writable folder found for upload")


# ── Jinja filters & globals ──────────────────────────────────────────────────
app.jinja_env.filters["score_color"] = score_color
app.jinja_env.filters["cos_rad"] = lambda deg: math.cos(math.radians(float(deg)))
app.jinja_env.filters["sin_rad"] = lambda deg: math.sin(math.radians(float(deg)))

@app.context_processor
def inject_user():
    u = current_user()
    if u:
        class _U:
            is_authenticated = True
            def __init__(self, d):
                self.__dict__.update(d)
                self.full_name = d.get("full_name", "")
                self.role      = d.get("role", "user")
                self.email     = d.get("email", "")
                self.github    = d.get("github", "")
                self.linkedin  = d.get("linkedin", "")
        return {"current_user": _U(u)}
    class _Anon:
        is_authenticated = False
        full_name = ""
        role      = ""
        email     = ""
        github    = ""
        linkedin  = ""
    return {"current_user": _Anon()}

# ── Auth decorators ──────────────────────────────────────────────────────────
def login_required(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        if not session.get("user_id"):
            flash("Please sign in to continue.", "warning")
            return redirect(url_for("login", next=request.path))
        u = current_user()
        if u is None:
            session.clear()
            flash("Session expired. Please sign in again.", "warning")
            return redirect(url_for("login", next=request.path))
        return f(*args, **kwargs)
    return decorated

def recruiter_required(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        if not session.get("user_id"):
            flash("Please sign in.", "warning")
            return redirect(url_for("login"))
        u = current_user()
        if u is None:
            session.clear()
            flash("Session expired. Please sign in again.", "warning")
            return redirect(url_for("login"))
        if u.get("role") != "recruiter":
            flash("Recruiter access required.", "error")
            return redirect(url_for("dashboard"))
        return f(*args, **kwargs)
    return decorated

def user_required(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        if not session.get("user_id"):
            flash("Please sign in.", "warning")
            return redirect(url_for("login", next=request.path))
        u = current_user()
        if u is None:
            session.clear()
            flash("Session expired. Please sign in again.", "warning")
            return redirect(url_for("login"))
        if u.get("role") != "user":
            flash("This area is for job seekers.", "error")
            return redirect(url_for("recruiter"))
        return f(*args, **kwargs)
    return decorated

# ════════════════════════════════════════════════════════════════════════════
# AUTH ROUTES
# ════════════════════════════════════════════════════════════════════════════

@app.route("/signup", methods=["GET", "POST"])
def signup():
    if session.get("user_id"):
        return redirect(url_for("dashboard"))

    if request.method == "POST":
        full_name = request.form.get("full_name", "").strip()
        email     = request.form.get("email", "").strip().lower()
        password  = request.form.get("password", "")
        confirm   = request.form.get("confirm_password", "")
        role      = request.form.get("role", "user")
        github    = request.form.get("github", "").strip()
        linkedin  = request.form.get("linkedin", "").strip()

        errors = []
        if len(full_name) < 2:
            errors.append("Please enter your full name.")
        if not re.match(r"^[^\s@]+@[^\s@]+\.[^\s@]+$", email):
            errors.append("Enter a valid email address.")
        if len(password) < 8:
            errors.append("Password must be at least 8 characters.")
        if password != confirm:
            errors.append("Passwords do not match.")
        if role not in ("user", "recruiter"):
            role = "user"

        with _users_lock:
            _load_users()
            if email in USERS:
                errors.append("An account with this email already exists.")

            if errors:
                for e in errors:
                    flash(e, "error")
                return render_template("signup.html")

            USERS[email] = {
                "id":            email,
                "full_name":     full_name,
                "email":         email,
                "password_hash": generate_password_hash(password),
                "role":          role,
                "github":        github,
                "linkedin":      linkedin,
                "created_at":    datetime.utcnow(),
            }
            _save_users()

        flash("Account created! Welcome aboard.", "success")
        session.clear()
        session["user_id"] = email
        session.permanent = True
        if role == "recruiter":
            return redirect(url_for("recruiter"))
        return redirect(url_for("dashboard"))

    return render_template("signup.html")


@app.route("/login", methods=["GET", "POST"])
def login():
    if session.get("user_id"):
        u = current_user()
        if u and u.get("role") == "recruiter":
            return redirect(url_for("recruiter"))
        return redirect(url_for("dashboard"))

    if request.method == "POST":
        email    = request.form.get("email", "").strip().lower()
        password = request.form.get("password", "")

        with _users_lock:
            _load_users()
            user = USERS.get(email)

        if user and check_password_hash(user["password_hash"], password):
            session.clear()
            session["user_id"] = email
            session.permanent  = True   # always permanent — HF proxy drops non-permanent cookies on redirect
            session.modified   = True
            flash(f"Welcome back, {user['full_name'].split()[0]}!", "success")
            if user["role"] == "recruiter":
                return redirect(url_for("recruiter"))
            next_url = request.form.get("next") or request.args.get("next", "")
            if next_url and next_url.startswith("/") and not next_url.startswith("//"):
                return redirect(next_url)
            return redirect(url_for("dashboard"))
        else:
            flash("Incorrect email or password.", "error")

    return render_template("login.html")


@app.route("/logout")
def logout():
    session.clear()
    flash("You've been signed out.", "info")
    return redirect(url_for("home"))


# ════════════════════════════════════════════════════════════════════════════
# DASHBOARD & PROFILE
# ════════════════════════════════════════════════════════════════════════════

@app.route("/dashboard")
@login_required
def dashboard():
    u = current_user()
    if not u:
        session.clear()
        return redirect(url_for("login"))
    if u.get("role") == "recruiter":
        return redirect(url_for("recruiter"))
    user_analyses = u.get("analyses_count", session.get("analyses_count", 0))
    user_best     = u.get("best_score",     session.get("best_score", None))
    return render_template("dashboard.html",
        greeting=get_greeting(),
        analyses_count=user_analyses,
        best_score=user_best,
        interview_sets=0,
        recent_analyses=[],
        global_stats=_load_stats(),
    )


@app.route("/profile", methods=["GET", "POST"])
@login_required
def profile():
    u = current_user()
    if not u:
        return redirect(url_for("login"))

    if request.method == "POST":
        full_name    = request.form.get("full_name", "").strip()
        email        = request.form.get("email", "").strip().lower()
        github       = request.form.get("github", "").strip()
        linkedin     = request.form.get("linkedin", "").strip()
        new_password = request.form.get("new_password", "")

        errors = []
        if len(full_name) < 2:
            errors.append("Name is too short.")
        if not re.match(r"^[^\s@]+@[^\s@]+\.[^\s@]+$", email):
            errors.append("Invalid email address.")
        if email != u["email"] and email in USERS:
            errors.append("That email is already in use.")

        if errors:
            for e in errors:
                flash(e, "error")
            return render_template("profile.html")

        old_email = u["email"]
        u["full_name"] = full_name
        u["github"]    = github
        u["linkedin"]  = linkedin
        if new_password:
            if len(new_password) < 8:
                flash("New password must be at least 8 characters.", "error")
                return render_template("profile.html")
            u["password_hash"] = generate_password_hash(new_password)
        if email != old_email:
            USERS[email] = u
            del USERS[old_email]
            session["user_id"] = email
            u["email"] = email
        else:
            USERS[old_email] = u

        _save_users()
        flash("Profile updated!", "success")
        return redirect(url_for("profile"))

    return render_template("profile.html")


# ════════════════════════════════════════════════════════════════════════════
# PUBLIC ROUTES
# ════════════════════════════════════════════════════════════════════════════

@app.route("/")
def index():
    """Landing page — always shown to everyone. Authenticated users also see it."""
    return render_template("index.html", global_stats=_load_stats())


@app.route("/home")
def home():
    return redirect(url_for("index"))


# ── Free public ATS score check ──────────────────────────────────────────────

@app.route("/check-ats", methods=["GET"])
def quick_ats():
    return render_template("quick_ats.html",
        categories=["None"] + list(JOB_ROLES.keys()),
        result=None,
        global_stats=_load_stats())


@app.route("/check-ats", methods=["POST"])
def quick_ats_post():
    categories = ["None"] + list(JOB_ROLES.keys())

    if "resume" not in request.files or request.files["resume"].filename == "":
        flash("Please select a file.", "error")
        return redirect(url_for("quick_ats"))

    f = request.files["resume"]
    if not allowed_file(f.filename):
        flash("Only PDF or DOCX files are accepted.", "error")
        return redirect(url_for("quick_ats"))

    fp = None
    try:
        fp = _save_upload(f)
        with open(fp, "rb") as fh:
            class _W:
                def __init__(self, x, n): self._f = x; self.name = n
                def read(self):           return self._f.read()
                def seek(self, *a):       return self._f.seek(*a)
            parsed = parser.parse(_W(fh, secure_filename(f.filename)))
            text = parsed.get("raw_text", "").strip()
    except Exception as e:
        flash(f"Could not read file: {e}", "error")
        return redirect(url_for("quick_ats"))
    finally:
        if fp and os.path.exists(fp):
            try: os.remove(fp)
            except Exception: pass

    if not text:
        flash("No text could be extracted from your file.", "error")
        return redirect(url_for("quick_ats"))

    cat  = request.form.get("category", "None")
    role = request.form.get("role", "None")
    if cat == "None" or role == "None":
        ri = {"description": "General analysis.", "required_skills": []}
    else:
        ri = JOB_ROLES.get(cat, {}).get(role, {"description": "", "required_skills": []})

    try:
        analysis = analyzer.analyze_resume({"raw_text": text}, ri)
    except Exception as e:
        flash(f"Analysis error: {e}", "error")
        return redirect(url_for("quick_ats"))

    ml = get_ats_ml(text)
    if ml is not None:
        analysis["ats_score"] = ml

    score = int(analysis.get("ats_score", 0))
    session["analyses_count"] = session.get("analyses_count", 0) + 1
    best = session.get("best_score")
    if best is None or score > best:
        session["best_score"] = score
    session.modified = True
    _increment_stats(score)

    uid = session.get("user_id")
    if uid and uid in USERS:
        u = USERS[uid]
        u["analyses_count"] = u.get("analyses_count", 0) + 1
        prev_best = u.get("best_score")
        if prev_best is None or score > prev_best:
            u["best_score"] = score
        _save_users()

    return render_template("quick_ats.html",
        categories=categories,
        result=analysis,
        global_stats=_load_stats())


@app.route("/about")
def about():
    return render_template("about.html")


# ════════════════════════════════════════════════════════════════════════════
# ATS ANALYZER (login required)
# ════════════════════════════════════════════════════════════════════════════

@app.route("/analyze", methods=["GET"])
@user_required
def analyze():
    return render_template("analyze.html",
        categories=["None"] + list(JOB_ROLES.keys()))


@app.route("/get_roles")
def get_roles():
    cat = request.args.get("category", "None")
    return jsonify([] if cat == "None" or cat not in JOB_ROLES
                   else list(JOB_ROLES[cat].keys()))


@app.route("/analyze", methods=["POST"])
@user_required
def analyze_post():
    if "resume" not in request.files or request.files["resume"].filename == "":
        flash("Please select a file.", "error")
        return redirect(url_for("analyze"))
    f = request.files["resume"]
    if not allowed_file(f.filename):
        flash("Only PDF/DOCX accepted.", "error")
        return redirect(url_for("analyze"))

    fp = None
    try:
        fp = _save_upload(f)
        with open(fp, "rb") as fh:
            class _W:
                def __init__(self, x, n): self._f = x; self.name = n
                def read(self):           return self._f.read()
                def seek(self, *a):       return self._f.seek(*a)
            parsed = parser.parse(_W(fh, secure_filename(f.filename)))
            text   = parsed.get("raw_text", "").strip()
    except Exception as e:
        flash(f"Read error: {e}", "error")
        return redirect(url_for("analyze"))
    finally:
        if fp and os.path.exists(fp):
            try: os.remove(fp)
            except Exception: pass

    if not text:
        flash("No text extracted from the file.", "error")
        return redirect(url_for("analyze"))

    cat  = request.form.get("category", "None")
    role = request.form.get("role", "None")
    if cat == "None" or role == "None":
        ri = {"description": "General analysis.", "required_skills": []}
        dr = "General Analysis"; dc = "General"
    else:
        ri = JOB_ROLES.get(cat, {}).get(role, {"description": "", "required_skills": []})
        dr = role; dc = cat

    try:
        analysis = analyzer.analyze_resume({"raw_text": text}, ri)
    except Exception as e:
        flash(f"Analysis error: {e}", "error")
        return redirect(url_for("analyze"))

    ml = get_ats_ml(text)
    if ml is not None:
        analysis["ats_score"] = ml

    if analysis.get("document_type") not in ("resume", None):
        flash(f"Looks like a '{analysis['document_type']}' — upload a resume.", "warning")
        return redirect(url_for("analyze"))

    score = int(analysis.get("ats_score", 0))
    session["analyses_count"] = session.get("analyses_count", 0) + 1
    best = session.get("best_score")
    if best is None or score > best:
        session["best_score"] = score
    session.modified = True
    _increment_stats(score)

    uid = session.get("user_id")
    if uid and uid in USERS:
        u = USERS[uid]
        u["analyses_count"] = u.get("analyses_count", 0) + 1
        prev_best = u.get("best_score")
        if prev_best is None or score > prev_best:
            u["best_score"] = score
        _save_users()

    return render_template("result.html", analysis=analysis,
                           role=dr, category=dc, role_info=ri)


# ════════════════════════════════════════════════════════════════════════════
# RESUME BUILDER
# ════════════════════════════════════════════════════════════════════════════

@app.route("/builder", methods=["GET"])
@user_required
def builder():
    return render_template("builder.html")


@app.route("/builder", methods=["POST"])
@user_required
def builder_post():
    if not DOCX_LOADED:
        flash("python-docx is not installed.", "error")
        return redirect(url_for("builder"))

    data      = request.form
    full_name = data.get("full_name", "").strip()
    email_val = data.get("email", "").strip()
    phone     = data.get("phone", "").strip()
    location  = data.get("location", "").strip()
    linkedin  = data.get("linkedin", "").strip()
    portfolio = data.get("portfolio", "").strip()
    summary   = data.get("summary", "").strip()

    if not full_name or not email_val:
        flash("Full name and email are required.", "error")
        return redirect(url_for("builder"))

    def _parse_json(key, default):
        try:
            return json.loads(data.get(key, default))
        except Exception:
            return json.loads(default)

    experiences = _parse_json("experiences_json", "[]")
    educations  = _parse_json("education_json",   "[]")
    projects    = _parse_json("projects_json",    "[]")
    skills_data = _parse_json("skills_json",      "{}")

    doc = DocxDocument()
    for sec in doc.sections:
        sec.top_margin    = Inches(0.75)
        sec.bottom_margin = Inches(0.75)
        sec.left_margin   = Inches(1.0)
        sec.right_margin  = Inches(1.0)

    def _heading(text):
        p   = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text.upper())
        run.bold = True
        run.font.size = Pt(11)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after  = Pt(2)
        pPr  = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bot  = OxmlElement("w:bottom")
        bot.set(qn("w:val"),   "single")
        bot.set(qn("w:sz"),    "6")
        bot.set(qn("w:space"), "1")
        bot.set(qn("w:color"), "2196F3")
        pBdr.append(bot)
        pPr.append(pBdr)

    def _bullet(text):
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(text).font.size = Pt(10)
        p.paragraph_format.space_after = Pt(1)

    def _body(text):
        p = doc.add_paragraph(text)
        if p.runs:
            p.runs[0].font.size = Pt(10)
        p.paragraph_format.space_after = Pt(2)

    np_ = doc.add_paragraph()
    np_.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nr  = np_.add_run(full_name)
    nr.bold = True; nr.font.size = Pt(20)

    contact = " | ".join(x for x in [email_val, phone, location, linkedin, portfolio] if x)
    cp = doc.add_paragraph(contact)
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if cp.runs:
        cp.runs[0].font.size = Pt(9)
    cp.paragraph_format.space_after = Pt(4)

    if summary:
        _heading("Professional Summary")
        _body(summary)

    if experiences:
        _heading("Work Experience")
        for exp in experiences:
            p = doc.add_paragraph()
            p.add_run(exp.get("position", "")).bold = True
            p.add_run(f"  —  {exp.get('company', '')}")
            dr = f"{exp.get('start_date', '')} – {exp.get('end_date', 'Present')}"
            tr = p.add_run(f"\t{dr}"); tr.font.size = Pt(9)
            p.paragraph_format.space_after = Pt(1)
            if exp.get("description"): _body(exp["description"])
            for r in exp.get("responsibilities", []):
                if r: _bullet(r)
            for a in exp.get("achievements", []):
                if a: _bullet(f"✓ {a}")

    if educations:
        _heading("Education")
        for edu in educations:
            p = doc.add_paragraph()
            p.add_run(f"{edu.get('degree', '')} in {edu.get('field', '')}").bold = True
            p.add_run(f"  —  {edu.get('school', '')}")
            meta = "  |  ".join(x for x in [
                edu.get("graduation_date", ""),
                f"GPA: {edu.get('gpa')}" if edu.get("gpa") else ""
            ] if x)
            if meta:
                tr = p.add_run(f"\t{meta}"); tr.font.size = Pt(9)
            for a in edu.get("achievements", []):
                if a: _bullet(a)

    if projects:
        _heading("Projects")
        for proj in projects:
            p = doc.add_paragraph()
            p.add_run(proj.get("name", "")).bold = True
            if proj.get("technologies"):
                tr = p.add_run(f"  |  {proj['technologies']}"); tr.font.size = Pt(9)
            if proj.get("link"):
                tr = p.add_run(f"  [{proj['link']}]"); tr.font.size = Pt(9)
            if proj.get("description"): _body(proj["description"])
            for r in proj.get("responsibilities", []):
                if r: _bullet(r)
            for a in proj.get("achievements", []):
                if a: _bullet(f"✓ {a}")

    if skills_data:
        _heading("Skills")
        for label, key in [("Technical", "technical"), ("Soft Skills", "soft"),
                            ("Languages", "languages"), ("Tools & Tech", "tools")]:
            vals = skills_data.get(key, [])
            if vals:
                p = doc.add_paragraph()
                p.add_run(f"{label}: ").bold = True
                p.add_run(", ".join(vals)).font.size = Pt(10)
                p.paragraph_format.space_after = Pt(1)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    filename = f"{full_name.replace(' ', '_')}_resume.docx"
    return send_file(
        buf,
        as_attachment=True,
        download_name=filename,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


# ════════════════════════════════════════════════════════════════════════════
# AI INTERVIEW QUESTIONS
# ════════════════════════════════════════════════════════════════════════════

ENCOURAGING_MSGS = [
    "All the best for your interview!",
    "Hope these help you prepare well!",
    "Good luck with your interview preparation!",
    "You're going to ace this!",
    "Keep up the great work!",
    "Believe in yourself and your abilities!",
    "Remember, preparation is key to success!",
    "You've got this! Go show them what you're made of!",
    "Stay confident and positive!",
    "Every question is a step towards your goal!",
]


@app.route("/ai-questions", methods=["GET", "POST"])
@user_required
def ai_questions():
    questions   = []
    prompt      = ""
    skills_str  = ""
    error       = None
    warning     = None
    encouraging = None

    if request.method == "POST":
        prompt = request.form.get("prompt", "").strip()
        try:
            from utils.ai_engine import get_questions_from_prompt
            result     = get_questions_from_prompt(prompt)
            questions  = result.get("questions", [])
            skills_str = result.get("skills_str", "")
            warning    = result.get("warning")
            error      = result.get("error")
            if questions:
                encouraging = random.choice(ENCOURAGING_MSGS)
        except Exception as exc:
            logger.error(f"ai_questions error: {exc}")
            error = f"Question generation failed: {exc}"

    return render_template("ai_questions.html",
        prompt=prompt, questions=questions, skills_str=skills_str,
        error=error, warning=warning, encouraging=encouraging)


# ════════════════════════════════════════════════════════════════════════════
# MOCK JOB SEARCH
# ════════════════════════════════════════════════════════════════════════════

_MOCK_JOBS = [
    {"title":"Software Engineer","company":"Infosys","location":"Hyderabad, India","salary":"₹6,00,000 – ₹12,00,000","url":"https://www.infosys.com/careers/","desc":"Design and develop scalable software solutions. Work with Java, Python, and cloud technologies.","date":"2025-04-01","tags":["software","engineer","java","python","developer"]},
    {"title":"Full Stack Developer","company":"TCS","location":"Bangalore, India","salary":"₹7,00,000 – ₹14,00,000","url":"https://www.tcs.com/careers","desc":"Build end-to-end web applications using React, Node.js, and SQL databases in an agile environment.","date":"2025-04-03","tags":["full stack","fullstack","react","node","javascript","developer","web"]},
    {"title":"Data Analyst","company":"Wipro","location":"Pune, India","salary":"₹5,00,000 – ₹10,00,000","url":"https://careers.wipro.com/","desc":"Analyse large datasets, create dashboards, and generate actionable insights using Python and SQL.","date":"2025-04-05","tags":["data","analyst","python","sql","analytics"]},
    {"title":"Machine Learning Engineer","company":"HCL Technologies","location":"Noida, India","salary":"₹10,00,000 – ₹20,00,000","url":"https://www.hcltech.com/careers","desc":"Develop and deploy ML models using TensorFlow and scikit-learn. Experience with NLP is a plus.","date":"2025-04-02","tags":["machine learning","ml","ai","data science","python","tensorflow"]},
    {"title":"Android Developer","company":"Tech Mahindra","location":"Hyderabad, India","salary":"₹6,00,000 – ₹11,00,000","url":"https://careers.techmahindra.com/","desc":"Build native Android apps using Java and Kotlin. Collaborate with UI/UX teams for seamless user experience.","date":"2025-04-04","tags":["android","mobile","java","kotlin","developer"]},
    {"title":"DevOps Engineer","company":"Capgemini","location":"Mumbai, India","salary":"₹8,00,000 – ₹16,00,000","url":"https://www.capgemini.com/careers/","desc":"Manage CI/CD pipelines, Docker containers, and AWS infrastructure. Jenkins and Kubernetes experience preferred.","date":"2025-04-06","tags":["devops","aws","docker","kubernetes","cloud","engineer"]},
    {"title":"Backend Developer","company":"Mphasis","location":"Bangalore, India","salary":"₹7,00,000 – ₹13,00,000","url":"https://careers.mphasis.com/","desc":"Develop RESTful APIs and microservices using Python/Django or Java/Spring Boot with MySQL databases.","date":"2025-04-07","tags":["backend","python","django","java","spring","api","developer"]},
    {"title":"Data Scientist","company":"IBM India","location":"Bangalore, India","salary":"₹12,00,000 – ₹22,00,000","url":"https://www.ibm.com/in-en/employment/","desc":"Work on advanced ML and statistical models to solve complex business problems at scale.","date":"2025-04-10","tags":["data science","data scientist","ml","python","statistics","analytics"]},
    {"title":"Software Engineer (Fresher)","company":"Accenture","location":"Multiple Cities, India","salary":"₹4,00,000 – ₹7,00,000","url":"https://www.accenture.com/in-en/careers","desc":"Entry-level role for fresh graduates. Training provided in Java, Python, and cloud platforms.","date":"2025-04-15","tags":["fresher","software","engineer","entry level","trainee","java","python"]},
    {"title":"QA Engineer","company":"Zoho","location":"Chennai, India","salary":"₹5,00,000 – ₹10,00,000","url":"https://www.zoho.com/careers.html","desc":"Design and execute test plans, automate test cases using Selenium, and ensure product quality.","date":"2025-04-19","tags":["qa","quality","testing","automation","selenium","engineer"]},
]

def _mock_job_search(query, location):
    if not query:
        return []
    q_lower = query.lower()
    q_words = re.split(r"[\s,/]+", q_lower)
    scored = []
    for job in _MOCK_JOBS:
        tag_str = " ".join(job["tags"])
        title_l = job["title"].lower()
        score = 0
        for w in q_words:
            if len(w) < 2: continue
            if w in tag_str: score += 2
            elif any(w in t for t in job["tags"]): score += 1
            if w in title_l: score += 3
        if q_lower in tag_str or q_lower in title_l: score += 5
        if score > 0: scored.append((score, job))
    if scored:
        scored.sort(key=lambda x: x[0], reverse=True)
        results = [j for _, j in scored]
    else:
        results = list(_MOCK_JOBS)
    if location:
        loc_lower = location.lower()
        filtered = [j for j in results if loc_lower in j["location"].lower()]
        if filtered: results = filtered
    return results


@app.route("/job-search", methods=["GET", "POST"])
@user_required
def job_search():
    results    = []
    query      = ""
    location_q = ""
    error      = None

    if request.method == "POST":
        query      = request.form.get("query", "").strip()
        location_q = request.form.get("location", "").strip()

        if not query:
            error = "Please enter a job title or keyword."
        else:
            app_id  = os.environ.get("ADZUNA_APP_ID",  "")
            app_key = os.environ.get("ADZUNA_APP_KEY", "")

            if app_id and app_key:
                try:
                    import requests as req_lib
                    from urllib.parse import quote as url_quote
                    country = "in"
                    api_url = (
                        f"https://api.adzuna.com/v1/api/jobs/{country}/search/1"
                        f"?app_id={app_id}&app_key={app_key}"
                        f"&results_per_page=20"
                        f"&what={url_quote(query)}"
                        f"&where={url_quote(location_q)}"
                        f"&content-type=application/json"
                    )
                    resp = req_lib.get(api_url, timeout=8)
                    if resp.ok:
                        for job in resp.json().get("results", []):
                            results.append({
                                "title":    job.get("title", ""),
                                "company":  job.get("company",  {}).get("display_name", "N/A"),
                                "location": job.get("location", {}).get("display_name", "N/A"),
                                "salary":   _fmt_salary(job),
                                "url":      job.get("redirect_url", "#"),
                                "desc":     (job.get("description", "")[:200] + "…"),
                                "date":     job.get("created", "")[:10],
                            })
                    else:
                        error = f"Job API error ({resp.status_code}). Showing curated listings."
                        results = _mock_job_search(query, location_q)
                except Exception as exc:
                    error = f"Could not fetch jobs: {exc}"
                    results = _mock_job_search(query, location_q)
            else:
                results = _mock_job_search(query, location_q)

    return render_template("job_search.html",
        results=results, query=query, location_q=location_q, error=error)


# ════════════════════════════════════════════════════════════════════════════
# FEEDBACK
# ════════════════════════════════════════════════════════════════════════════

@app.route("/feedback", methods=["GET"])
@login_required
def feedback():
    avg_rating = 0
    if _feedback_store:
        avg_rating = round(
            sum(f["rating"] for f in _feedback_store) / len(_feedback_store), 1
        )
    dist = {1: 0, 2: 0, 3: 0, 4: 0, 5: 0}
    for f in _feedback_store:
        dist[f["rating"]] = dist.get(f["rating"], 0) + 1

    return render_template("feedback.html",
        feedbacks=_feedback_store,
        avg_rating=avg_rating,
        rating_dist=dist,
        total=len(_feedback_store),
    )


@app.route("/feedback", methods=["POST"])
@login_required
def feedback_post():
    u       = current_user()
    name    = request.form.get("name",           u.get("full_name", "")).strip()
    email_v = request.form.get("feedback_email", u.get("email", "")).strip()
    rating  = max(1, min(5, int(request.form.get("rating", 3))))
    feature = request.form.get("feature", "").strip()
    comment = request.form.get("comment", "").strip()

    if not comment:
        flash("Please write a comment before submitting.", "error")
        return redirect(url_for("feedback"))

    _feedback_store.append({
        "name":    name,
        "email":   email_v,
        "rating":  rating,
        "feature": feature,
        "comment": comment,
        "date":    datetime.utcnow().strftime("%d %b %Y"),
    })
    _save_feedback()
    flash("Thank you for your feedback! 🙏", "success")
    return redirect(url_for("feedback"))


# ════════════════════════════════════════════════════════════════════════════
# RECRUITER ROUTES
# ════════════════════════════════════════════════════════════════════════════

EDU_KW = {
    "Schooling":      ["10th","high school","ssc","secondary school","matriculation","matric","class x","class 10"],
    "Intermediate":   ["12th","intermediate","hsc","senior secondary","puc","class xii","class 12","higher secondary","plus two","+2"],
    "Bachelors":      ["b.tech","b.e.","b.sc","b.a.","b.com","btech","bsc","bachelor","bachelors","undergraduate","b.eng","b.s.","honours"],
    "Masters":        ["m.tech","m.e.","m.sc","m.b.a","mtech","msc","mba","master","masters","postgraduate","m.eng","m.s.","pg diploma","pgdm"],
    "PhD":            ["ph.d","ph.d.","doctorate","doctoral","d.phil"],
    "Certifications": ["diploma","certificate","certification","certified","coursera","udemy","nptel","aws certified","google certified","microsoft certified"],
}
EDU_KW_BOUNDARY = {
    "Schooling":      {"ssc","matric"},
    "Intermediate":   {"hsc","puc"},
    "Bachelors":      {"btech","bsc","bachelor","bachelors","undergraduate"},
    "Masters":        {"mtech","msc","mba","master","masters","postgraduate"},
    "PhD":            {"doctorate","doctoral"},
    "Certifications": {"diploma","certificate","certification","certified"},
}
TECH_SKILLS = {
    "Prog":   ["c","c++","java","python","javascript","ruby","go","kotlin","swift","php","scala"],
    "Web":    ["html","css","node.js","react","django","flask","angular","vue.js","bootstrap"],
    "DB":     ["mysql","mongodb","postgresql","sqlite","oracle","nosql"],
    "FW":     ["pandas","numpy","tensorflow","pytorch","scikit-learn","spring"],
    "Cloud":  ["aws","google cloud","azure","docker","kubernetes"],
    "Tools":  ["git","github","jenkins","jira","linux","android studio"],
    "Mobile": ["android","ios","flutter","react native"],
}
SCORE_PAT = re.compile(r"(?i)(?:gpa|cgpa|percentage|marks)[:\s-]*([\d\.]+(?:/[\d\.]+)?%?)")


def _extract_text_r(fbytes, ftype, fname=""):
    try:
        fname_lower = fname.lower()
        if "pdf" in ftype or fname_lower.endswith(".pdf"):
            if PDFMINER_LOADED:
                return pdfminer_extract(io.BytesIO(fbytes))
            import PyPDF2
            r = PyPDF2.PdfReader(io.BytesIO(fbytes))
            return "\n".join(p.extract_text() or "" for p in r.pages)
        if "word" in ftype or fname_lower.endswith(".docx"):
            if DOCX_LOADED:
                doc = DocxDocument(io.BytesIO(fbytes))
                return "\n".join(p.text for p in doc.paragraphs)
        import PyPDF2
        r = PyPDF2.PdfReader(io.BytesIO(fbytes))
        return "\n".join(p.extract_text() or "" for p in r.pages)
    except Exception:
        return ""


def _is_resume(text, from_csv=False):
    if from_csv: return True
    tl  = text.lower()
    scs = sum(s in tl for s in ["education","experience","skills","projects","internships"])
    kws = [k for v in EDU_KW.values() for k in v] + [s for v in TECH_SKILLS.values() for s in v]
    return scs >= 2 or sum(k in tl for k in kws) >= 5


def _norm_score(s_str, _level):
    try:
        s_str = s_str.lower()
        if "%" in s_str:
            m = re.search(r"(\d+\.?\d*)%", s_str); return float(m.group(1)) if m else None
        m = re.search(r"(\d+\.?\d*)\s*/\s*(\d+\.?\d*)", s_str)
        if m:
            s, mx = float(m.group(1)), float(m.group(2)); return (s / mx) * 100
        m = re.search(r"(\d+\.?\d*)", s_str)
        if m:
            s = float(m.group(1))
            if s <= 4:  return (s / 4)  * 100
            if s <= 10: return (s / 10) * 100
            return s if s <= 100 else None
    except Exception:
        return None


def _edu_kw_match(ll, cat, kws):
    boundary_set = EDU_KW_BOUNDARY.get(cat, set())
    for kw in kws:
        if kw in boundary_set:
            if re.search(rf"\b{re.escape(kw)}\b", ll): return True
        else:
            if kw in ll: return True
    return False


def _extract_edu(text):
    lines = [l.strip() for l in text.split("\n")]
    raw   = []
    i     = 0
    while i < len(lines):
        ll = lines[i].lower()
        if not ll: i += 1; continue
        deg = next(
            (cat for cat, kws in EDU_KW.items() if _edu_kw_match(ll, cat, kws)),
            ""
        )
        if deg:
            window = " ".join(lines[i : i + 5])
            sm = SCORE_PAT.search(window)
            sr = sm.group(0) if sm else "N/A"
            ns = _norm_score(sr, deg) if sm else None
            ym = re.search(r"\b(?:19|20)\d{2}\b", window)
            raw.append({"degree": deg, "score": sr, "normalized_score": ns,
                        "year": ym.group(0) if ym else ""})
        i += 1
    best = {}
    for entry in raw:
        d = entry["degree"]
        existing = best.get(d)
        if existing is None:
            best[d] = entry
        elif existing["normalized_score"] is None and entry["normalized_score"] is not None:
            best[d] = entry
    order = list(EDU_KW.keys())
    return sorted(best.values(), key=lambda e: order.index(e["degree"]) if e["degree"] in order else 99)


def _extract_skills_r(text, required):
    skills = set(); tl = text.lower()
    for sl in TECH_SKILLS.values():
        for s in sl:
            if re.search(rf"\b{re.escape(s)}\b", tl): skills.add(s)
    for s in required:
        ss = s if isinstance(s, str) else " or ".join(s)
        if re.search(rf"\b{re.escape(ss)}\b", tl):
            skills.update(s if isinstance(s, list) else [s])
    return skills


def _parse_criteria(q):
    edu, skills, thresh = {}, [], 0
    pat = re.compile(
        r"(schooling|intermediate|bachelors?|masters?|phd|certifications?)\s*"
        r"(?:above|with)?\s*(\d+\.?\d*)\s*(%|gpa)?", re.I)
    for m in pat.finditer(q):
        lv, sc, un = m.group(1).capitalize(), float(m.group(2)), (m.group(3) or "").lower()
        pct = (sc if un == "%" else
               ((sc / 4) * 100 if un == "gpa" and sc <= 4 else
                (sc / 10) * 100 if un == "gpa" else sc))
        edu[lv] = max(0, min(pct, 100))
    pool = {s for v in TECH_SKILLS.values() for s in v}
    for w in q.lower().split():
        if w in pool: skills.append(w)
    am = re.search(r"ats\s*(?:score)?\s*(?:above)?\s*(\d+)", q, re.I)
    if am: thresh = int(am.group(1))
    return {"edu": edu, "skills": list(set(skills)), "thresh": thresh}


def _score_candidate(name, text, crit):
    edu    = _extract_edu(text)
    skills = _extract_skills_r(text, crit["skills"])
    ml     = get_ats_ml(text)
    if ml is None:
        r  = analyzer.analyze_resume({"raw_text": text}, {"required_skills": crit["skills"]})
        ml = r.get("ats_score", 50)
    ats    = round(float(ml), 1)
    m_edu  = (not crit["edu"]) or all(
        any(e["degree"].lower() == lv.lower() and
            e["normalized_score"] and e["normalized_score"] >= req for e in edu)
        for lv, req in crit["edu"].items())
    m_sk   = (not crit["skills"]) or all(
        (any(s in [x.lower() for x in skills] for s in sk) if isinstance(sk, list)
         else sk.lower() in [s.lower() for s in skills])
        for sk in crit["skills"])
    m_ats  = ats >= crit["thresh"] if crit["thresh"] else True
    em     = (sum(
        any(e["degree"].lower() == lv.lower() and
            e["normalized_score"] and e["normalized_score"] >= req for e in edu)
        for lv, req in crit["edu"].items()) / len(crit["edu"])) if crit["edu"] else 1.0
    sm_val = (sum(
        (any(s in [x.lower() for x in skills] for s in sk)
         if isinstance(sk, list) else sk.lower() in [s.lower() for s in skills])
        for sk in crit["skills"]) / len(crit["skills"])) if crit["skills"] else 1.0
    overall = round((em + sm_val + (1 if m_ats else 0)) / 3 * 100, 1)
    return {
        "name":    name,
        "ats":     ats,
        "match":   overall,
        "edu":     "; ".join(f"{e['degree']} ({e['score']})" for e in edu) or "N/A",
        "skills":  ", ".join(sorted(skills)) or "N/A",
        "edu_met": m_edu,
        "sk_met":  m_sk,
        "ats_met": m_ats,
        "ok":      (m_edu and m_sk and m_ats) or overall >= 50,
    }


@app.route("/recruiter")
@recruiter_required
def recruiter():
    return render_template("recruiter.html", candidates=session.get("candidates", []))


def _proc_csv(fbytes, fname, crit):
    import csv as csv_mod
    rows_out = []
    try:
        text_content = fbytes.decode("utf-8", errors="replace")
        reader = csv_mod.DictReader(io.StringIO(text_content))
        fieldnames = reader.fieldnames or []
        name_col = next(
            (c for c in fieldnames if c.strip().lower() in ("name", "candidate", "candidate_name", "full_name", "applicant")),
            None
        )
        for i, row in enumerate(reader, start=1):
            if name_col:
                cand_name = row.pop(name_col, "").strip() or f"Candidate_{i}"
            else:
                cand_name = f"Row_{i}"
            resume_text = " ".join(str(v) for v in row.values() if v and str(v).strip())
            if not resume_text.strip():
                rows_out.append((None, f"{fname}:row{i}"))
                continue
            rows_out.append((_score_candidate(cand_name, resume_text, crit), None))
    except Exception as exc:
        logger.warning(f"CSV parse error for {fname}: {exc}")
    return rows_out


@app.route("/recruiter/screen", methods=["POST"])
@recruiter_required
def rec_screen():
    files = request.files.getlist("resumes")
    query = request.form.get("criteria", "").strip()
    if not files or all(f.filename == "" for f in files):
        flash("Upload at least one resume.", "error")
        return redirect(url_for("recruiter"))

    invalid = [f.filename for f in files if f.filename and not allowed_recruiter_file(f.filename)]
    if invalid:
        flash(f"Unsupported file type(s): {', '.join(invalid)}. Only PDF, DOCX and CSV are accepted.", "error")
        return redirect(url_for("recruiter"))

    crit = _parse_criteria(query) if query else {"edu": {}, "skills": [], "thresh": 0}
    results = []; skipped = []

    def _proc(uf):
        nm  = secure_filename(uf.filename)
        fb  = uf.read()
        ft  = uf.content_type or ""
        if nm.lower().endswith(".csv") or "csv" in ft:
            return "CSV", _proc_csv(fb, nm, crit)
        text = _extract_text_r(fb, ft, nm)
        if not text or not _is_resume(text): return None, nm
        return _score_candidate(nm, text, crit), None

    with ThreadPoolExecutor(max_workers=4) as ex:
        for outcome in ex.map(_proc, [f for f in files if f.filename != ""]):
            if outcome[0] == "CSV":
                for r, s in outcome[1]:
                    if r: results.append(r)
                    if s: skipped.append(s)
            else:
                r, s = outcome
                if r: results.append(r)
                if s: skipped.append(s)

    results.sort(key=lambda x: x["match"], reverse=True)
    session["candidates"] = results
    if skipped: flash(f"Skipped: {', '.join(skipped[:10])}{'…' if len(skipped)>10 else ''}", "warning")
    flash(f"Processed {len(results)} candidate(s) successfully.", "success")
    return redirect(url_for("recruiter"))


@app.route("/recruiter/export")
@recruiter_required
def rec_export():
    import csv
    cands = session.get("candidates", [])
    if not cands:
        flash("Nothing to export.", "warning")
        return redirect(url_for("recruiter"))
    si = io.StringIO()
    w  = csv.DictWriter(si, fieldnames=cands[0].keys())
    w.writeheader(); w.writerows(cands)
    return Response(si.getvalue(), mimetype="text/csv",
                    headers={"Content-Disposition": "attachment;filename=shortlisted.csv"})


# ════════════════════════════════════════════════════════════════════════════
# ERROR HANDLERS
# ════════════════════════════════════════════════════════════════════════════

@app.errorhandler(404)
def e404(e):
    return render_template("error.html", code=404, message="Page not found."), 404

@app.errorhandler(413)
def e413(e):
    flash("File too large. Regular uploads: max 10 MB.", "error")
    return redirect(url_for("index"))

@app.errorhandler(500)
def e500(e):
    logger.error(f"500 error: {e}")
    return render_template("error.html", code=500, message="Something went wrong on our end."), 500


# ════════════════════════════════════════════════════════════════════════════
# AI ENHANCEMENT ROUTES
# ════════════════════════════════════════════════════════════════════════════

from utils.ai_engine import (
    get_smart_feedback,
    get_skill_gaps,
    get_interview_questions,
    get_rewrite_suggestions,
    get_role_optimization,
)


def _parse_and_extract(route_name):
    """Shared logic: validate upload, save to tmp, parse text. Returns (text, error_response)."""
    if "resume" not in request.files or request.files["resume"].filename == "":
        flash("Please upload a resume.", "error")
        return None, redirect(url_for(route_name))
    f = request.files["resume"]
    if not allowed_file(f.filename):
        flash("Only PDF/DOCX accepted.", "error")
        return None, redirect(url_for(route_name))
    fp = None
    try:
        fp = _save_upload(f)
        with open(fp, "rb") as fh:
            class _W:
                def __init__(self, x, n): self._f = x; self.name = n
                def read(self):           return self._f.read()
                def seek(self, *a):       return self._f.seek(*a)
            parsed = parser.parse(_W(fh, secure_filename(f.filename)))
            text = parsed.get("raw_text", "").strip()
    except Exception as e:
        flash(f"Could not read file: {e}", "error")
        return None, redirect(url_for(route_name))
    finally:
        if fp and os.path.exists(fp):
            try: os.remove(fp)
            except Exception: pass
    if not text:
        flash("No text could be extracted from the file.", "error")
        return None, redirect(url_for(route_name))
    return text, None


@app.route("/ai-feedback", methods=["GET", "POST"])
@user_required
def ai_feedback():
    result = None
    error  = None

    if request.method == "POST":
        text, err_resp = _parse_and_extract("ai_feedback")
        if err_resp: return err_resp

        cat  = request.form.get("category", "None")
        role = request.form.get("role", "None")
        ri   = {} if cat == "None" or role == "None" else JOB_ROLES.get(cat, {}).get(role, {})
        required = ri.get("required_skills", [])
        job_role = role if role != "None" else ""

        try:
            result = get_smart_feedback(text, job_role, required)
            result["resume_text"] = text[:500]
        except Exception as e:
            logger.error(f"AI feedback error: {e}")
            error = f"Analysis error: {e}"

    return render_template("ai_feedback.html", result=result, error=error,
                           categories=list(JOB_ROLES.keys()))


@app.route("/skill-gap", methods=["GET", "POST"])
@user_required
def skill_gap():
    result = None
    error  = None

    if request.method == "POST":
        text, err_resp = _parse_and_extract("skill_gap")
        if err_resp: return err_resp

        cat  = request.form.get("category", "None")
        role = request.form.get("role", "None")
        ri   = {} if cat == "None" or role == "None" else JOB_ROLES.get(cat, {}).get(role, {})
        required = ri.get("required_skills", [])
        job_role = role if role != "None" else "Software Engineer"

        try:
            result = get_skill_gaps(text, job_role, required, cat)
        except Exception as e:
            logger.error(f"Skill gap error: {e}")
            error = f"Analysis error: {e}"

    return render_template("skill_gap.html", result=result, error=error,
                           categories=list(JOB_ROLES.keys()))


@app.route("/smart-questions", methods=["GET", "POST"])
@user_required
def smart_questions():
    result = None
    error  = None

    if request.method == "POST":
        cat   = request.form.get("category", "None")
        role  = request.form.get("role", "None")
        level = request.form.get("level", "mid")
        num   = min(int(request.form.get("num_questions", 15)), 50)

        ri = {} if cat == "None" or role == "None" else JOB_ROLES.get(cat, {}).get(role, {})
        skills = ri.get("required_skills", [])
        job_role = role if role != "None" else "Software Engineer"

        try:
            result = get_interview_questions(job_role, skills, level, num)
        except Exception as e:
            logger.error(f"Smart questions error: {e}")
            error = f"Generation error: {e}"

    return render_template("smart_questions.html", result=result, error=error,
                           categories=list(JOB_ROLES.keys()))


@app.route("/rewrite", methods=["GET", "POST"])
@user_required
def rewrite():
    result = None
    error  = None

    if request.method == "POST":
        text, err_resp = _parse_and_extract("rewrite")
        if err_resp: return err_resp

        cat  = request.form.get("category", "None")
        role = request.form.get("role", "None")
        job_role = role if role != "None" else ""

        try:
            result = get_rewrite_suggestions(text, job_role)
        except Exception as e:
            logger.error(f"Rewrite error: {e}")
            error = f"Analysis error: {e}"

    return render_template("rewrite.html", result=result, error=error,
                           categories=list(JOB_ROLES.keys()))


@app.route("/role-optimizer", methods=["GET", "POST"])
@user_required
def role_optimizer():
    result = None
    error  = None

    if request.method == "POST":
        text, err_resp = _parse_and_extract("role_optimizer")
        if err_resp: return err_resp

        cat  = request.form.get("category", "None")
        role = request.form.get("role", "None")
        ri   = {} if cat == "None" or role == "None" else JOB_ROLES.get(cat, {}).get(role, {})
        required = ri.get("required_skills", [])
        job_role = role if role != "None" else "Software Engineer"

        try:
            result = get_role_optimization(text, job_role, required, cat)
        except Exception as e:
            logger.error(f"Role optimizer error: {e}")
            error = f"Analysis error: {e}"

    return render_template("role_optimizer.html", result=result, error=error,
                           categories=list(JOB_ROLES.keys()))


# ════════════════════════════════════════════════════════════════════════════
if __name__ == "__main__":
    port = int(os.environ.get("PORT", 7860))
    app.run(debug=False, host="0.0.0.0", port=port)